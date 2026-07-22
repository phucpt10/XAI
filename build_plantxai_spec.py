from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


REFERENCE = r"D:\ResearchCode\XAI\Template_DacTa_ResearchSoftware.docx"
OUTPUT = r"D:\ResearchCode\XAI\PlantXAI-Stability_Research_Software_Specification_v0.9_Draft.docx"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF5FB"
GRAY = "F2F2F2"
MID_GRAY = "D9E1F2"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_body(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in (("Title", 25, BLUE), ("Subtitle", 12, "666666"), ("Heading 1", 16, BLUE), ("Heading 2", 12.5, BLUE), ("Heading 3", 11, BLUE)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(10 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    if "Code Block" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(8.5)
        code.paragraph_format.left_indent = Inches(0.22)
        code.paragraph_format.right_indent = Inches(0.12)
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run(run, color=BLUE, bold=True)
    return p


def add_para(doc, text="", bold_prefix=None, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run(r2, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run(r)
    return p


def add_code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), GRAY)
    pPr.append(shd)
    r = p.add_run(text)
    set_run(r, size=8.5, font="Consolas")
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9360])
    repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    set_run(r, bold=True, color=BLUE)
    r = p.add_run(text)
    set_run(r)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None, header_fill=BLUE, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = str(value)
        shade(cell, header_fill)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(1)
            for run in p.runs:
                set_run(run, size=font_size, bold=True, color="FFFFFF")
    repeat_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for run in p.runs:
                    set_run(run, size=font_size)
    if widths:
        set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_module(doc, name, purpose, inputs, outputs, logic, edge_cases):
    add_heading(doc, name, 2)
    add_para(doc, purpose, bold_prefix="Purpose: ")
    add_para(doc, "Inputs:", bold_prefix="Inputs:")
    add_table(doc, ["Name", "Type", "Source", "Description"], inputs, [1800, 1800, 2100, 3660], font_size=8.5)
    add_para(doc, "Outputs:", bold_prefix="Outputs:")
    add_table(doc, ["Name", "Type", "Destination", "Description"], outputs, [1800, 1800, 2100, 3660], font_size=8.5)
    add_para(doc, "Processing logic:", bold_prefix="Processing logic:")
    add_code(doc, logic)
    add_para(doc, "Edge cases and error handling:", bold_prefix="Edge cases and error handling:")
    add_table(doc, ["Condition", "Required handling"], edge_cases, [3600, 5760], font_size=8.5)


def build():
    doc = Document(REFERENCE)
    clear_body(doc)
    configure_styles(doc)
    section = doc.sections[0]
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(82)
    r = p.add_run("PLANTXAI-STABILITY")
    set_run(r, size=25, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research Software Design and Specification")
    set_run(r, size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Developer Edition — Scientific Protocol v0.9 Draft")
    set_run(r, size=12, italic=True, color="666666")
    doc.add_paragraph()
    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(meta, [2850, 6510])
    metadata = [
        ("Document purpose", "Executable research software specification"),
        ("Scientific status", "Draft; G0A=PASS; G0B_PROTOCOL_FREEZE_READY=BLOCKED"),
        ("Protocol version", "v0.9"),
        ("Document version", "1.0 English draft"),
        ("Scope", "PlantVillage tomato disease classification and XAI stability"),
        ("Audience", "Research software architect, developer, reviewer and AI pair programmer"),
    ]
    for i, (k, v) in enumerate(metadata):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        shade(meta.cell(i, 0), LIGHT_BLUE)
        for j, cell in enumerate(meta.rows[i].cells):
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, size=9, bold=j == 0, color=BLUE if j == 0 else None)
    repeat_header(meta.rows[0])
    doc.add_paragraph()
    add_callout(doc, "Draft status:", "This specification freezes the intended scientific behavior before implementation. It contains no fabricated dataset counts, model results, p-values or confidence intervals.")
    doc.add_page_break()

    # Design philosophy and contents
    add_heading(doc, "DESIGN PHILOSOPHY", 1)
    add_para(doc, "PlantXAI-Stability is research software rather than a conventional production application. Its quality target is not only that the code runs, but that every scientific conclusion is traceable, reproducible and auditable.")
    add_table(doc, ["Dimension", "Conventional software", "Research software"], [
        ("Objective", "Feature works and scales", "Runs, reproduces and verifies scientific claims"),
        ("Quality", "Bug-free behavior", "Traceable artifacts and scientific invariants"),
        ("Change policy", "Iterative product change", "Versioned and frozen protocol changes"),
        ("Primary users", "End users and developers", "Researchers, reviewers and developers"),
        ("Output", "Feature or service", "Auditable evidence bundle"),
    ], [1800, 3600, 3960], font_size=8.5)
    add_callout(doc, "Core principle:", "Freeze the scientific intent before writing the experiment code. A pipeline that runs but silently changes the population, target class or statistical unit can produce invalid results.")
    add_heading(doc, "DOCUMENT STRUCTURE", 1)
    add_table(doc, ["Part", "Name", "Purpose"], [
        ("A", "Research Context", "Why the study exists and which questions it answers"),
        ("B", "Scope and Boundaries", "What is included, excluded and claimable"),
        ("C", "System Architecture", "Layers, modules, data flow and technology"),
        ("D", "Module Specifications", "Typed interfaces, algorithms and error handling"),
        ("E", "Data Contracts", "Stable schemas and artifact formats"),
        ("F", "Business Rules", "Implementation rules that must not be violated"),
        ("G", "Scientific Invariants", "Mathematical and methodological truths"),
        ("H", "Quality Gates", "Acceptance criteria from setup to release"),
        ("I", "Testing Strategy", "Unit, integration, scientific and reproducibility tests"),
        ("J", "Handoff Protocol", "Traceability, task breakdown and review"),
    ], [900, 2700, 5760], font_size=8.5)
    doc.add_page_break()

    # A
    add_heading(doc, "PART A: RESEARCH CONTEXT", 1)
    add_para(doc, "Purpose: define the scientific problem, questions and contribution before implementation.")
    add_heading(doc, "A1. Official title", 2)
    add_para(doc, "PlantXAI-Stability: Controlled Evaluation of Prediction Robustness and Explanation Stability in Plant Disease Classification")
    add_heading(doc, "A2. Problem statement", 2)
    add_para(doc, "A classifier may preserve its predicted disease label after a controlled image perturbation while changing the visual evidence highlighted by an XAI method. Prediction robustness and explanation stability are therefore distinct properties. PlantXAI-Stability evaluates both properties jointly, while preventing explanation comparisons when the prediction target itself has changed.")
    add_heading(doc, "A3. Research questions", 2)
    add_table(doc, ["RQ", "Question", "Required evidence", "Primary artifacts"], [
        ("RQ1", "How robust are ResNet50 and EfficientNet-B0 to controlled image transformations?", "Original/transformed class, confidence, consistency and severity trends", "prediction_results.parquet; Table 3"),
        ("RQ2", "How do Grad-CAM, Grad-CAM++ and Score-CAM differ in explanation stability?", "Paired valid heatmaps, SSIM, Pearson, cosine and uncertainty", "stability_metrics.parquet; Tables 4–5"),
        ("RQ3", "How is prediction robustness related to explanation stability?", "Joint records; confidence change versus heatmap similarity in consistent subset", "joint_results.parquet; relationship figures"),
    ], [700, 3100, 3100, 2460], font_size=8)
    add_heading(doc, "A4. Expected contributions", 2)
    for item in [
        "A controlled protocol for evaluating prediction and explanation behavior together.",
        "A paired benchmark of three CAM methods on two CNN backbones.",
        "An analysis of transformation type and severity effects.",
        "Evidence about cases where prediction remains stable but visual evidence changes.",
        "A reproducible research system with manifests, hashes, frozen splits, exclusion reasons and artifact lineage.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "A5. Audience and publication boundary", 2)
    add_para(doc, "Primary users are explainable-AI researchers, plant-vision researchers, research software engineers, reviewers and developers implementing the protocol. The project supports a methods or empirical research paper; it is not a production diagnosis service.")
    doc.add_page_break()

    # B
    add_heading(doc, "PART B: SCOPE AND BOUNDARIES", 1)
    add_heading(doc, "B1. In scope", 2)
    for item in [
        "PlantVillage primary confirmatory population: five tomato classes.",
        "ResNet50 and EfficientNet-B0 image classifiers.",
        "Grad-CAM, Grad-CAM++ and Score-CAM.",
        "Rotation, brightness, Gaussian noise and Gaussian blur; three pilot-approved severities each.",
        "Prediction consistency, confidence change, SSIM, Pearson correlation and cosine similarity.",
        "Leaf-aware bootstrap confidence intervals, paired Wilcoxon tests, rank-biserial effect sizes and Holm correction.",
        "Machine-readable per-sample records, publication tables/figures and reproducibility manifests.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "B2. Out of scope", 2)
    for item in [
        "A new XAI algorithm.",
        "Faithfulness or causal explanation evaluation unless separately specified.",
        "Direct claims about field deployment, other crops or other datasets.",
        "Model selection using official test performance.",
        "Manual editing of scientific results in paper tables or figures.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "B3. Assumptions and pending decisions", 2)
    for item in [
        "Protocol v0.9 remains draft until dataset, leaf identity, splits, severity, checkpoint and target layer are approved.",
        "The source dataset provides a valid leaf-level grouping key; absence of leaf_id is a blocking condition.",
        "Training and validation configuration values are provisional until pilot evidence is reviewed.",
        "A numerical tolerance, rather than exact GPU equality, may be required for floating-point reproducibility.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "B4. Permitted claims", 2)
    add_callout(doc, "Permitted:", "On the five tomato classes of PlantVillage, the framework measures how prediction consistency and CAM similarity change under controlled transformations.", fill="E2F0D9")
    add_callout(doc, "Not permitted:", "The framework proves that one CAM method is universally better, faithful or ready for field diagnosis.", fill="FCE4D6")
    add_heading(doc, "B5. Limitations", 2)
    for item in [
        "PlantVillage may contain background, acquisition and collection-process shortcuts.",
        "Controlled transformations do not represent all real-world distribution shifts.",
        "Stability is evaluated only for prediction-consistent pairs.",
        "A stable explanation is not automatically a faithful or clinically correct explanation.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # C
    add_heading(doc, "PART C: SYSTEM ARCHITECTURE", 1)
    add_heading(doc, "C1. Layered architecture", 2)
    add_code(doc, "CLI / experiment scripts\n    ↓\nApplication orchestration: Trainer, JointEvaluator\n    ↓\nScientific domain services: Prediction, XAI, Alignment, Metrics, Statistics\n    ↓\nTyped interfaces and immutable data contracts\n    ↓\nInfrastructure: PyTorch, filesystem, Parquet, plotting\n\nCross-cutting: Configuration, Provenance, Logging, Tests and Quality Gates")
    add_para(doc, "Dependencies flow downward. Scientific packages must not import CLI code or notebooks. Configuration and provenance are inputs to orchestration and domain services, not hidden global state.")
    add_heading(doc, "C2. Scientific data flow", 2)
    add_code(doc, "Frozen manifest → original inference → controlled transformation → transformed inference\n        → consistency gate\n          ├─ inconsistent: prediction robustness record + exclusion\n          └─ consistent: same-target CAM pair → inverse alignment → quality gate\n                         → SSIM/Pearson/Cosine → joint record\n        → leaf-aware statistics → tables, figures and reproducibility bundle")
    add_heading(doc, "C3. Package map", 2)
    add_table(doc, ["Package", "Responsibility", "Must not contain"], [
        ("config", "Schema, validation, merge, resolved-config hash", "Inference, dataset I/O, plotting"),
        ("data", "Manifest, hash, frozen splits, Dataset/DataLoader", "Training or XAI logic"),
        ("transformations", "Four transformations, severity and inverse metadata", "Model or statistics decisions"),
        ("models", "Wrappers, training, checkpoints and inference", "Paper plotting"),
        ("xai", "Grad-CAM family, heatmap generation and normalization", "Dataset splitting or hypothesis tests"),
        ("evaluation", "Robustness, consistency, alignment, stability and joint records", "Hard-coded experiment parameters"),
        ("statistics", "Bootstrap, Wilcoxon, effect size and Holm", "Image processing"),
        ("visualization", "Tables, figures and qualitative examples", "Recomputing scientific metrics"),
        ("provenance", "Run identity, hashes, environment and lineage", "Scientific decision logic"),
    ], [1900, 4560, 2900], font_size=8)
    add_heading(doc, "C4. Technology stack", 2)
    add_table(doc, ["Layer", "Technology", "Version/status", "Selection rationale"], [
        ("Core ML", "PyTorch", ">=2.1 (planned)", "Backbone training, inference and hooks"),
        ("Data", "Hugging Face Datasets / Parquet", "Pinned revision", "Schema-aware data access and immutable records"),
        ("XAI", "Captum or validated adapters", ">=0.6 (planned)", "CAM implementations under a common interface"),
        ("Statistics", "SciPy / statsmodels", "Pinned environment", "Non-parametric tests and correction"),
        ("Validation", "pytest, Ruff, MyPy", "Pinned environment", "Software and scientific quality gates"),
    ], [1700, 2500, 1900, 3260], font_size=8)
    add_heading(doc, "C5. Directory structure", 2)
    add_code(doc, "project/\n├── configs/              # versioned protocol and resolved configs\n├── src/plantxai_stability/\n│   ├── config/\n│   ├── data/\n│   ├── transformations/\n│   ├── models/\n│   ├── xai/\n│   ├── evaluation/\n│   ├── statistics/\n│   ├── visualization/\n│   └── provenance/\n├── experiments/          # thin orchestration entry points\n├── tests/                 # unit, integration, scientific, reproducibility\n├── data/                  # raw, manifests, splits, interim cache\n├── outputs/runs/          # immutable run-specific artifacts\n└── docs/                  # decisions, protocol and reports")
    add_heading(doc, "C6. Central application service", 2)
    add_para(doc, "JointEvaluator coordinates the scientific sequence, enforces the consistency policy and passes the original predicted class to both CAM calls. It must not implement model internals, image transformations, metrics, statistical tests or plotting itself.")
    add_code(doc, "JointEvaluator\n├── DatasetManifestBuilder / SplitManager\n├── TransformationPipeline\n├── InferenceEngine\n├── PredictionMetrics\n├── XAIGenerator\n├── HeatmapAligner\n├── ExplanationMetrics\n├── RelationshipAnalyzer\n├── StatisticalAnalysis\n└── ArtifactGenerator")
    doc.add_page_break()

    # D
    add_heading(doc, "PART D: MODULE SPECIFICATIONS", 1)
    add_para(doc, "Each module has a narrow responsibility, typed contracts and explicit failure behavior. All scientific parameters are passed from the resolved configuration.")
    add_module(doc, "D1. ConfigLoader", "Loads, merges and validates versioned YAML configuration; returns an immutable resolved configuration and its hash.", [
        ("config_paths", "list[path]", "CLI", "Dataset, model, transformation, metric and statistics YAML files"),
        ("schema", "JSON Schema", "config package", "Structural and typed validation rules"),
    ], [
        ("resolved_config", "Frozen mapping", "all services", "Validated values with defaults resolved"),
        ("config_hash", "str", "provenance", "SHA-256 identity for the resolved configuration"),
    ], "resolved = merge_sources(config_paths)\nvalidate(resolved, schema)\nreject_unknown_keys(resolved)\nreturn freeze(resolved), sha256(canonical_json(resolved))", [
        ("Unknown key", "Fail closed with a configuration error."),
        ("Missing required value", "Fail the G0 gate; do not apply hidden defaults."),
        ("Conflicting source values", "Require explicit resolution and record it."),
    ])
    add_module(doc, "D2. DatasetManifestBuilder and SplitManager", "Creates stable image identity, audits integrity and reads approved frozen splits without silently re-splitting data.", [
        ("dataset_revision", "Pinned identifier", "dataset receipt", "Immutable source revision"),
        ("class_policy", "list[str]", "protocol", "Five approved tomato classes"),
        ("leaf_id", "str", "dataset metadata", "Group key for split and bootstrap"),
    ], [
        ("manifest", "Parquet", "data/manifests", "Identity, class, leaf, dimensions and hashes"),
        ("split_report", "JSON", "quality gates", "Overlap, counts and split hashes"),
    ], "read source receipt\ncanonicalize RGB pixels\ncompute canonical_rgb_sha256\ncreate deterministic sample_id\n audit duplicates and leaf overlap\nload or write approved frozen splits", [
        ("Missing leaf_id", "Hard failure; never invent a grouping key."),
        ("Duplicate across splits", "Fail data gate and emit a duplicate report."),
        ("Corrupted image", "Exclude only with a documented reason code."),
    ])
    add_module(doc, "D3. TransformationPipeline", "Applies four controlled perturbations at three pilot-approved severities and returns exact transformation metadata.", [
        ("pixel_tensor", "float tensor", "dataset adapter", "Canonical RGB pixels in [0,1]"),
        ("scenario", "TransformationConfig", "config", "Type, severity and parameter policy"),
        ("seed_context", "SeedContext", "provenance", "Global seed, sample_id and scenario_id"),
    ], [
        ("transformed_pixels", "float tensor", "inference", "Deterministic transformed input"),
        ("transformation_record", "record", "evaluation", "Actual parameter, seed, inverse and mask metadata"),
    ], "seed = derive(global_seed, sample_id, scenario_id)\nparams = resolve_params(scenario, seed)\ntransformed = apply(pixel_tensor, params)\nrecord exact params, inverse metadata and valid mask\nreturn transformed, record", [
        ("Unsupported scenario", "Reject configuration before execution."),
        ("Invalid severity", "Fail validation; do not clamp silently."),
        ("Rotation border", "Record fill policy and produce valid-overlap mask."),
    ])
    add_module(doc, "D4. ModelWrapper and InferenceEngine", "Provides a common interface for ResNet50 and EfficientNet-B0, including checkpoint identity and deterministic inference.", [
        ("model_config", "ModelConfig", "config", "Backbone, classifier and preprocessing"),
        ("checkpoint", "file reference", "approved artifact", "Validation-selected checkpoint"),
        ("model_tensor", "float tensor", "dataset adapter", "Resized and ImageNet-normalized input"),
    ], [
        ("prediction_record", "Parquet row", "evaluation", "Logits, class, probability and confidence"),
        ("target_layer", "layer reference", "XAI", "Runtime-validated convolutional layer"),
    ], "load checkpoint and verify hash\nset eval mode and no_grad for inference\nlogits = model(model_tensor)\nprobabilities = softmax(logits)\nreturn class, confidence, logits and provenance", [
        ("Checkpoint hash mismatch", "Block execution and report provenance failure."),
        ("Wrong output dimension", "Fail model gate; expected five classes."),
        ("Missing target layer", "Fail target-layer validation before CAM."),
    ])
    add_module(doc, "D5. XAIGenerator", "Generates Grad-CAM, Grad-CAM++ and Score-CAM heatmaps for an explicitly supplied target class and target layer.", [
        ("model", "ModelWrapper", "models", "Loaded model and hooks interface"),
        ("input_tensor", "float tensor", "inference", "Original or transformed model input"),
        ("target_class", "int", "JointEvaluator", "Original predicted class"),
        ("method", "enum", "config", "One of three CAM methods"),
    ], [
        ("heatmap", "float array", "alignment/metrics", "Two-dimensional CAM map"),
        ("cam_record", "record", "provenance", "Method, target, layer, shape and checksum"),
    ], "validate target class and target layer\nattach hooks\nrun selected CAM algorithm\nresize to input resolution\nremove hooks\nreturn heatmap and metadata", [
        ("Hook leak", "Fail the test and remove hooks in finally blocks."),
        ("Constant heatmap", "Return quality flag; do not assign stability."),
        ("Score-CAM memory pressure", "Chunk and resume without changing the scientific target."),
    ])
    add_module(doc, "D6. JointEvaluator and HeatmapAligner", "Pairs original and transformed predictions, enforces consistency, aligns geometric heatmaps and creates joint records.", [
        ("original_predictions", "Parquet", "inference cache", "One record per model and sample"),
        ("transformed_predictions", "Parquet", "inference", "One record per model, sample and scenario"),
        ("transformation_record", "record", "transformation", "Inverse transform and valid mask"),
    ], [
        ("joint_record", "Parquet row", "statistics", "Prediction, CAM and stability linkage"),
        ("exclusion_record", "Parquet row", "audit", "Documented reason for invalid pair"),
    ], "join by run_id, model_id, sample_id and scenario_id\ncompute consistency\nif inconsistent: record exclusion and stop CAM\nelse: call CAM twice with original predicted class\nalign transformed heatmap and write joint record", [
        ("Prediction changed", "Keep RQ1 record; exclude from official explanation stability."),
        ("Orphan or duplicate key", "Fail join audit and do not aggregate."),
        ("Empty valid mask", "Document alignment failure."),
    ])
    add_module(doc, "D7. ExplanationMetrics and StatisticalAnalysis", "Validates heatmaps, computes stability metrics and performs leaf-aware uncertainty and paired inference.", [
        ("aligned_heatmaps", "float arrays", "alignment", "Same shape, normalized and masked"),
        ("joint_records", "Parquet", "evaluation", "Per-sample valid records"),
        ("statistics_config", "StatisticsConfig", "config", "Bootstrap, test and correction policy"),
    ], [
        ("metric_records", "Parquet", "artifacts", "SSIM, Pearson, cosine per sample"),
        ("statistical_results", "Parquet", "reports", "CI, statistic, p, adjusted p and effect size"),
    ], "validate finite values and mask\ncompute SSIM(data_range=1.0), Pearson and cosine\nbootstrap by leaf_id\nintersect paired keys\nrun Wilcoxon\ncompute rank-biserial effect size\napply Holm by predeclared family", [
        ("Zero variance or norm", "Mark metric invalid with reason code."),
        ("Different paired sets", "Use paired intersection and report n_pairs."),
        ("Multiple comparisons", "Apply Holm to the declared family, never post hoc."),
    ])
    add_module(doc, "D8. ArtifactGenerator and RunManifestWriter", "Exports validated results and records the complete provenance needed to reproduce a run.", [
        ("validated_records", "Parquet", "statistics", "Per-sample summaries and test results"),
        ("provenance", "RunContext", "runtime", "Config, Git, environment and hashes"),
    ], [
        ("tables", "CSV/LaTeX", "publication", "Generated from one source of truth"),
        ("figures", "PNG/PDF", "publication", "300 DPI raster and vector where possible"),
        ("run_manifest", "JSON", "reproducibility", "Run, artifact and exclusion lineage"),
    ], "resolve output directory by run_id\nwrite atomic partitions and completed markers\nexport tables and figures\nhash artifacts\nwrite run_manifest and artifact_index", [
        ("Existing incompatible output", "Refuse to append; require a new run_id."),
        ("Table mismatch", "Fail artifact gate when CSV and LaTeX differ."),
        ("Dirty worktree", "Block official release unless explicitly permitted."),
    ])
    doc.add_page_break()

    # E
    add_heading(doc, "PART E: DATA CONTRACTS", 1)
    add_heading(doc, "E1. Immutable records", 2)
    add_code(doc, "@dataclass(frozen=True)\nclass SampleRecord:\n    sample_id: str\n    leaf_id: str\n    class_id: int\n    class_name: str\n    source_split: str\n    split: str\n    canonical_relative_path: str\n    canonical_rgb_sha256: str\n    width: int\n    height: int\n\n@dataclass(frozen=True)\nclass PredictionRecord:\n    run_id: str\n    model_id: str\n    sample_id: str\n    scenario_id: str\n    predicted_class: int\n    confidence: float\n    is_correct: bool\n    checkpoint_sha256: str\n\n@dataclass(frozen=True)\nclass TransformationRecord:\n    sample_id: str\n    scenario_id: str\n    seed: int\n    parameters_json: str\n    inverse_metadata_json: str\n    valid_mask_sha256: str\n\n@dataclass(frozen=True)\nclass JointRecord:\n    run_id: str\n    model_id: str\n    sample_id: str\n    leaf_id: str\n    scenario_id: str\n    xai_method: str\n    target_class: int\n    is_consistent: bool\n    ssim: float | None\n    pearson: float | None\n    cosine: float | None\n    exclusion_reason: str | None")
    add_heading(doc, "E2. File formats", 2)
    add_table(doc, ["Artifact", "Format", "Schema", "Immutable"], [
        ("dataset_manifest", "Parquet", "SampleRecord", "Yes after approval"),
        ("split manifests", "CSV/Parquet", "sample_id, leaf_id, class, split", "Yes after freeze"),
        ("prediction_results", "Parquet", "PredictionRecord", "Yes per run"),
        ("transformation_records", "Parquet", "TransformationRecord", "Yes per run"),
        ("stability_metrics", "Parquet", "JointRecord plus metrics", "Yes per run"),
        ("statistical_results", "Parquet", "Estimate, CI, p, adjusted p, effect", "Yes per run"),
        ("run_manifest", "JSON", "Provenance and artifact index", "Yes after finalization"),
    ], [2200, 1600, 3300, 2260], font_size=8)
    add_heading(doc, "E3. Naming and join conventions", 2)
    for item in [
        "Files use snake_case; columns use snake_case.",
        "Every result key includes run_id, model_id, sample_id and scenario_id; XAI records also include xai_method.",
        "No join may depend on filesystem order, row position or batch order.",
        "Each scientific configuration change creates a new run_id.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "E4. Versioning and output layout", 2)
    add_code(doc, "outputs/runs/<run_id>/\n├── resolved_config.yaml\n├── run_manifest.json\n├── predictions/\n├── heatmaps/\n├── metrics/\n├── statistics/\n├── tables/\n├── figures/\n├── reports/\n├── logs/\n├── status/\n└── artifact_index.json")
    add_para(doc, "Protocol version, resolved configuration hash, dataset revision, manifest hash, split hashes and checkpoint hashes are all required provenance fields.")
    doc.add_page_break()

    # F
    add_heading(doc, "PART F: BUSINESS RULES", 1)
    add_callout(doc, "Implementation policy:", "Violation of a business rule rejects the change request. Business rules are operational constraints; scientific invariants in Part G are mathematical or methodological truths.")
    add_table(doc, ["ID", "Rule", "Owner", "Test ID"], [
        ("BR-001", "Evaluate explanation stability only when original and transformed predicted classes are equal.", "JointEvaluator", "T-FIL-01"),
        ("BR-002", "Both CAM calls use the original predicted class as the target.", "XAIGenerator", "T-XAI-01"),
        ("BR-003", "Heatmaps are normalized to [0,1] with epsilon=1e-8 before comparison.", "Metrics", "T-XAI-02"),
        ("BR-004", "Geometric transformations provide inverse metadata and valid-overlap masks.", "Transform/Align", "T-ALN-01"),
        ("BR-005", "Transformation parameters are read from versioned YAML; no scientific hard-coding.", "Transform", "T-TRF-01"),
        ("BR-006", "sample_id is content/path canonicalization based, never row or batch position.", "Data", "T-DAT-01"),
        ("BR-007", "Official test is not used for model, severity or target-layer selection.", "Training", "T-TRN-01"),
        ("BR-008", "Statistical pairing uses explicit keys and the common valid intersection.", "Statistics", "T-STA-01"),
        ("BR-009", "Bootstrap respects leaf-level clustering when multiple images share a leaf.", "Statistics", "T-STA-02"),
        ("BR-010", "Holm families are declared before results are inspected.", "Statistics", "T-STA-03"),
        ("BR-011", "Constant, non-finite or invalid heatmaps receive an exclusion reason.", "Quality", "T-QA-01"),
        ("BR-012", "Weights, checkpoints, secrets and local AI configuration are not committed to Git.", "Release", "T-REL-01"),
    ], [1100, 4920, 1800, 1540], font_size=7.8)
    add_heading(doc, "F2. Anti-patterns", 2)
    for item in [
        "Hard-coded absolute paths or parameter values.",
        "Joining records by row position or DataLoader order.",
        "Using test metrics to choose a checkpoint, class, severity or target layer.",
        "Replacing a constant heatmap with a zero map and treating zero–zero similarity as valid.",
        "Dropping exclusions silently.",
        "Editing paper numbers manually.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()

    # G
    add_heading(doc, "PART G: SCIENTIFIC INVARIANTS", 1)
    add_para(doc, "The following statements must remain true regardless of implementation refactoring:")
    add_table(doc, ["ID", "Invariant", "Verification"], [
        ("SI-001", "Same seed, protocol and input produce the same stable identity and transformation parameters.", "Scientific test"),
        ("SI-002", "Official test membership is not used for any selection decision.", "AST/config audit"),
        ("SI-003", "Explanation stability is undefined for prediction-inconsistent pairs.", "Gate test"),
        ("SI-004", "Both CAM heatmaps in a pair use one common target class.", "Mocked XAI test"),
        ("SI-005", "Normalized heatmaps lie in [0,1] and invalid/constant maps are excluded.", "Assertion + quality test"),
        ("SI-006", "Inverse rotation followed by forward rotation approximately preserves spatial identity.", "Synthetic alignment test; SSIM threshold declared in protocol"),
        ("SI-007", "SSIM(image, image) is approximately 1.0 under the configured implementation.", "Sanity test"),
        ("SI-008", "Paired statistical methods use the same explicit paired keys.", "Pairing audit"),
        ("SI-009", "Holm adjusted p-values are monotone under the declared step-down procedure.", "Reference test"),
        ("SI-010", "Every expected pair is a valid result, documented exclusion or documented failure.", "Population reconciliation"),
    ], [1100, 5900, 2360], font_size=8)
    doc.add_page_break()

    # H
    add_heading(doc, "PART H: QUALITY GATES", 1)
    add_table(doc, ["Gate", "Stage", "Criteria", "Pass condition"], [
        ("G0", "Configuration", "Schema valid; resolved config and protocol state recorded", "Validator PASS"),
        ("G1", "Dataset", "Five classes, valid leaf identity, no duplicate/split leakage", "Audit PASS"),
        ("G2", "Baseline", "Two checkpoints selected using validation only", "Evidence and hash match"),
        ("G3", "Transformations", "12 scenarios deterministic and semantically piloted", "Transform tests PASS"),
        ("G4", "XAI", "Three CAM adapters and target layers runtime-validated", "Heatmap smoke PASS"),
        ("G5", "Joint evaluation", "Consistency gate, alignment and quality flags work", "Scientific tests PASS"),
        ("G6", "Statistics", "Leaf-aware bootstrap, paired tests, effect size and Holm", "Reference match"),
        ("G7", "Artifacts", "Tables/figures generated from validated records", "Cross-artifact check PASS"),
        ("G8", "Reproducibility", "Clean rerun and hashes/tolerances documented", "Rerun PASS"),
        ("G9", "Release", "Tests, docs, provenance and review complete", "CI/review green"),
    ], [900, 1900, 4760, 1800], font_size=7.8)
    add_heading(doc, "H2. Definition of Done", 2)
    for item in [
        "All public functions have type hints and Google-style docstrings.",
        "Unit, integration and scientific tests pass; target coverage is above 80% for core logic.",
        "No untracked scientific parameters or absolute paths.",
        "All exclusions and failures have reason codes.",
        "Tables and figures regenerate from machine-readable records.",
        "Run manifest, hashes and clean-rerun report are complete.",
        "Pull request is reviewed and traceability matrix is updated.",
    ]:
        add_bullet(doc, "[ ] " + item)
    doc.add_page_break()

    # I
    add_heading(doc, "PART I: TESTING STRATEGY", 1)
    add_heading(doc, "I1. Test pyramid", 2)
    add_code(doc, "                 /\\\n                /  \\       Scientific tests (5–10%)\n               / SCI \\      Invariants, gates and reference values\n              /------\\\n             /        \\     Integration tests (15–20%)\n            /   INT    \\    Module boundaries and contracts\n           /------------\\\n          /              \\  Unit tests (70–80%)\n         /      UNIT      \\ Pure functions, adapters and validators\n        /------------------\\")
    add_heading(doc, "I2. Test categories", 2)
    add_table(doc, ["Category", "Location", "Required coverage"], [
        ("Unit", "tests/unit/", "Functions, schemas, IDs, transforms, metric edge cases; core coverage >80%"),
        ("Integration", "tests/integration/", "Dataset → transform → inference; prediction → CAM → metric; statistics → artifacts"),
        ("Scientific", "tests/scientific/", "All SI rules, consistency gate, alignment, pairing and reference statistics"),
        ("Reproducibility", "tests/reproducibility/", "Two runs, hash comparison, numerical tolerance and clean environment"),
        ("Release/QA", "CI and scripts", "Lint, type checking, secret scan, artifact reconciliation and documentation"),
    ], [1700, 2200, 5460], font_size=8)
    add_heading(doc, "I3. Mandatory scientific test examples", 2)
    for item in [
        "A changed prediction produces a prediction record but no official CAM comparison.",
        "A rotation plus inverse rotation returns an image/heatmap close to identity under the predeclared tolerance.",
        "A constant heatmap is invalid even if a library would return a finite similarity.",
        "The same CAM target class is passed to original and transformed calls.",
        "A paired comparison rejects unequal sample keys rather than silently truncating lists.",
        "Bootstrap resampling at leaf level differs from image-level resampling in the expected contract test.",
        "Every expected pair is reconciled into result, exclusion or failure.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "I4. Reproducibility policy", 2)
    add_para(doc, "Exact determinism is required for IDs, splits, transformation metadata, filtering and aggregation. Numerical reproducibility for GPU operations is reported with an explicit tolerance; the system must not claim exact equality when hardware or library behavior prevents it.")
    doc.add_page_break()

    # J
    add_heading(doc, "PART J: HANDOFF PROTOCOL", 1)
    add_heading(doc, "J1. Developer handoff checklist", 2)
    for item in [
        "Specification and protocol state are complete.",
        "Decision Records exist for dataset, split, severity, target layer and statistical unit.",
        "Reference fixtures include duplicate, leakage, inconsistent prediction, constant heatmap and alignment cases.",
        "Golden files exist for schema, manifest, transformation record and statistical reference output.",
        "The developer has the acceptance criteria for each module and gate.",
        "No official experiment is run while G0B is blocked.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "J2. Traceability matrix", 2)
    add_table(doc, ["Research question", "Primary module", "Record/metric", "Test/gate", "Artifact"], [
        ("RQ1", "InferenceEngine + PredictionMetrics", "PredictionRecord; consistency; confidence delta", "G2/G5; T-PRED-*", "prediction_results; Table 3"),
        ("RQ2", "XAIGenerator + HeatmapAligner + ExplanationMetrics", "CAM pair; SSIM/Pearson/cosine", "G4/G5/G6; T-XAI-*", "stability_metrics; Tables 4–5"),
        ("RQ3", "JointEvaluator + RelationshipAnalyzer", "JointRecord; confidence versus stability", "G5/G6; T-RQ3-*", "relationship_analysis; figures"),
    ], [1200, 2500, 2800, 1500, 1360], font_size=7.5)
    add_heading(doc, "J3. Task breakdown template", 2)
    add_code(doc, "Task ID: TASK-[MODULE]-[N]\nModule: [module name]\nPriority: P0 / P1 / P2\nDependencies: [TASK-xxx]\nProtocol reference: [BR-xxx, SI-xxx]\n\nDescription:\n[Short implementation description]\n\nAcceptance criteria:\n- [ ] Typed interface and documented contract\n- [ ] Unit tests pass\n- [ ] Scientific tests pass where applicable\n- [ ] No hard-coded scientific parameter\n- [ ] Artifact/provenance behavior verified\n\nFiles to create or modify:\n- src/plantxai_stability/...\n- tests/...\n- docs/...")
    add_heading(doc, "J4. AI pair-programming prompt contract", 2)
    add_code(doc, "Role: Senior research software engineer\nProject: PlantXAI-Stability\nProtocol: v0.9 draft; do not freeze or change scientific intent\nTask: [specific task]\nBusiness rules: [BR references]\nScientific invariants: [SI references]\nData contracts: [record/schema references]\nRequired: type hints, docstrings, tests, provenance\nForbidden: fabricated results, test leakage, hard-coded parameters, silent exclusions")
    add_heading(doc, "J5. Review checklist", 2)
    for item in [
        "Code runs in a clean environment.",
        "Unit and integration tests pass.",
        "Scientific invariants are verified.",
        "Business rules are respected.",
        "No test leakage or silent exclusions exist.",
        "All results are traceable to run, config, data, model and code hashes.",
        "No claim exceeds the declared PlantVillage scope.",
    ]:
        add_bullet(doc, "[ ] " + item)
    doc.add_page_break()

    add_heading(doc, "CONCLUSION", 1)
    add_para(doc, "PlantXAI-Stability is a research experimentation framework for measuring whether stable model outputs are accompanied by stable visual evidence. Its central contract is:")
    add_callout(doc, "Scientific contract:", "Controlled dataset → frozen identity and splits → deterministic transformations → paired predictions → consistency filtering → same-target CAM pairs → geometric alignment → heatmap quality gate → per-sample stability metrics → leaf-aware statistics → reproducible evidence bundle.")
    add_para(doc, "The current document is a v0.9 draft. It is ready to guide implementation, but official scientific execution remains blocked until the dataset audit, leaf identity, split policy, severity pilot, checkpoint selection and target-layer validation are completed and approved.", italic=True)
    add_para(doc, "End of specification.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Footer
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("PlantXAI-Stability — Research Software Specification v0.9 Draft")
    set_run(r, size=8, color="666666")
    doc.core_properties.title = "PlantXAI-Stability Research Software Design and Specification"
    doc.core_properties.subject = "English developer edition scientific research software specification"
    doc.core_properties.author = "Research Software Architecture"
    doc.core_properties.comments = "Draft; no official experimental results included."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
