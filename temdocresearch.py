"""
Script sinh file Template_DacTa_ResearchSoftware.docx
Chạy: python generate_template.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    """Tô màu nền cho cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_code_block(doc, code_text):
    """Thêm code block với nền xám."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Nền xám
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)
    
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    
    doc.add_paragraph()  # spacing
    return table


def build_document():
    doc = Document()
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('TEMPLATE BẢNG THIẾT KẾ & ĐẶC TẢ\nPHẦN MỀM NGHIÊN CỨU', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Research Software Engineering Specification Template\n'
                           'Dành cho Research Software Architect → Developer / AI Pair Programming')
    run.font.size = Pt(12)
    run.italic = True
    
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Phiên bản: 1.0\n').bold = True
    meta.add_run('Ngày ban hành: 22/07/2026\n')
    meta.add_run('Áp dụng cho: PlantXAI-Stability và các đề tài tương tự')
    
    doc.add_page_break()
    
    # ===== TRIẾT LÝ =====
    add_heading_styled(doc, 'TRIẾT LÝ THIẾT KẾ', level=1)
    doc.add_paragraph(
        'Phần mềm nghiên cứu khác phần mềm thông thường ở 5 khía cạnh cốt lõi:'
    )
    add_table(doc,
        ['Khía cạnh', 'Software Engineering', 'Research Software Engineering'],
        [
            ['Mục tiêu', 'Chạy được, scale', 'Chạy + tái lập + kiểm chứng'],
            ['Người dùng', 'End-user, dev', 'Reviewer, cộng đồng KH'],
            ['Chất lượng', 'Bug-free', 'Traceable, reproducible'],
            ['Thay đổi', 'Agile, iterate', 'Frozen protocol, DR'],
            ['Kết quả', 'Feature shipped', 'Artifact có thể audit'],
        ])
    
    doc.add_paragraph(
        '→ Template này đóng băng ý định khoa học (frozen scientific intent) '
        'trước khi code được viết, tránh "code chạy nhưng kết quả sai".'
    )
    
    # ===== CẤU TRÚC 10 PHẦN =====
    add_heading_styled(doc, 'CẤU TRÚC TEMPLATE (10 PHẦN)', level=1)
    add_table(doc,
        ['Phần', 'Tên', 'Nội dung'],
        [
            ['A', 'Research Context', 'Tại sao làm? Làm gì?'],
            ['B', 'Scope & Boundaries', 'Làm đến đâu? Không làm gì?'],
            ['C', 'System Architecture', 'Cấu trúc tổng thể'],
            ['D', 'Module Specifications', 'Chi tiết từng module'],
            ['E', 'Data Contracts', 'Hợp đồng dữ liệu'],
            ['F', 'Business Rules', 'Quy tắc bất biến'],
            ['G', 'Scientific Invariants', 'Điều phải luôn đúng'],
            ['H', 'Quality Gates', 'Tiêu chí nghiệm thu'],
            ['I', 'Testing Strategy', 'Cách kiểm chứng'],
            ['J', 'Handoff Protocol', 'Bàn giao cho dev/AI'],
        ])
    
    doc.add_page_break()
    
    # ===== PHẦN A =====
    add_heading_styled(doc, 'PHẦN A: RESEARCH CONTEXT', level=1)
    doc.add_paragraph('Mục đích: Xác định "TẠI SAO" và "LÀM GÌ" trước khi code.')
    
    add_heading_styled(doc, 'A1. Tên đề tài', level=2)
    doc.add_paragraph('[Tên chính thức bằng tiếng Anh]')
    doc.add_paragraph('[Tên tiếng Việt - nếu có]')
    
    add_heading_styled(doc, 'A2. Vấn đề nghiên cứu (Problem Statement)', level=2)
    doc.add_paragraph('[Mô tả 2-3 câu: Vấn đề gì chưa được giải quyết? Tại sao quan trọng?]')
    
    add_heading_styled(doc, 'A3. Câu hỏi nghiên cứu (Research Questions)', level=2)
    add_table(doc,
        ['RQ ID', 'Câu hỏi', 'Module trả lời', 'Output artifact'],
        [
            ['RQ1', '...', '...', 'Table X, Fig Y'],
            ['RQ2', '...', '...', '...'],
            ['RQ3', '...', '...', '...'],
        ])
    
    add_heading_styled(doc, 'A4. Đóng góp dự kiến (Contributions)', level=2)
    doc.add_paragraph('• C1: [Framework/Method mới]')
    doc.add_paragraph('• C2: [Dataset/Benchmark]')
    doc.add_paragraph('• C3: [Empirical findings]')
    
    add_heading_styled(doc, 'A5. Target venue & audience', level=2)
    doc.add_paragraph('• Tạp chí/Hội nghị: [Tên, Q-ranking]')
    doc.add_paragraph('• Độc giả: [AI researchers, domain experts, ...]')
    
    doc.add_page_break()
    
    # ===== PHẦN B =====
    add_heading_styled(doc, 'PHẦN B: SCOPE & BOUNDARIES', level=1)
    
    add_heading_styled(doc, 'B1. IN SCOPE (Làm)', level=2)
    doc.add_paragraph('• Dataset: [Tên, nguồn, config]')
    doc.add_paragraph('• Models: [Liệt kê cụ thể]')
    doc.add_paragraph('• Methods: [Liệt kê cụ thể]')
    doc.add_paragraph('• Metrics: [Liệt kê cụ thể]')
    
    add_heading_styled(doc, 'B2. OUT OF SCOPE (Không làm)', level=2)
    doc.add_paragraph('• [Liệt kê rõ ràng để tránh scope creep]')
    doc.add_paragraph('• [Ví dụ: Không đánh giá faithfulness của XAI]')
    
    add_heading_styled(doc, 'B3. ASSUMPTIONS (Giả định)', level=2)
    doc.add_paragraph('• [Ví dụ: GPU memory >= 8GB]')
    doc.add_paragraph('• [Ví dụ: Dataset không có label noise]')
    
    add_heading_styled(doc, 'B4. INTENDED CLAIMS (Claim được phép)', level=2)
    doc.add_paragraph('✅ "Framework đánh giá được độ ổn định của XAI"')
    doc.add_paragraph('❌ "Framework chứng minh Grad-CAM tốt hơn Score-CAM" (overclaim)')
    
    add_heading_styled(doc, 'B5. LIMITATIONS (Hạn chế phải thừa nhận)', level=2)
    doc.add_paragraph('• [Liệt kê trung thực - reviewer sẽ tìm ra nếu giấu]')
    
    doc.add_page_break()
    
    # ===== PHẦN C =====
    add_heading_styled(doc, 'PHẦN C: SYSTEM ARCHITECTURE', level=1)
    
    add_heading_styled(doc, 'C1. High-level Architecture Diagram', level=2)
    doc.add_paragraph('[Chèn ASCII art hoặc link draw.io]')
    
    add_heading_styled(doc, 'C2. Data Flow Diagram', level=2)
    add_code_block(doc,
        '[Input] → [Module A] → [Module B] → [Output]\n'
        '              ↓            ↓\n'
        '          [Artifact 1] [Artifact 2]')
    
    add_heading_styled(doc, 'C3. Module Dependency Graph', level=2)
    doc.add_paragraph('[Module nào gọi module nào? Interface là gì?]')
    
    add_heading_styled(doc, 'C4. Technology Stack', level=2)
    add_table(doc,
        ['Layer', 'Technology', 'Version', 'Lý do chọn'],
        [
            ['Core ML', 'PyTorch', '≥2.1', '...'],
            ['Data', 'HuggingFace', '≥2.14', '...'],
            ['XAI', 'Captum', '≥0.6', '...'],
            ['Statistics', 'SciPy, statsmodels', '...', '...'],
        ])
    
    add_heading_styled(doc, 'C5. Directory Structure', level=2)
    add_code_block(doc,
        'project/\n'
        '├── configs/        # Frozen protocol configs\n'
        '├── src/            # Source code\n'
        '├── scripts/        # Utility scripts\n'
        '├── experiments/    # Main experiment runners\n'
        '├── tests/          # Unit + Scientific tests\n'
        '├── outputs/        # Gitignored artifacts\n'
        '└── docs/           # Decision records')
    
    doc.add_page_break()
    
    # ===== PHẦN D =====
    add_heading_styled(doc, 'PHẦN D: MODULE SPECIFICATIONS', level=1)
    doc.add_paragraph('Mỗi module cần một bảng đặc tả theo template sau:')
    
    add_heading_styled(doc, 'D[N]. [TÊN MODULE]', level=2)
    
    doc.add_paragraph('Mục đích:', style='List Bullet')
    doc.add_paragraph('[Mô tả 1-2 câu module này làm gì]')
    
    doc.add_paragraph('Inputs:', style='List Bullet')
    add_table(doc,
        ['Tên', 'Kiểu', 'Nguồn', 'Mô tả'],
        [
            ['input_1', 'torch.Tensor', 'Module trước', '...'],
            ['config', 'dict', 'YAML file', '...'],
        ])
    
    doc.add_paragraph('Outputs:', style='List Bullet')
    add_table(doc,
        ['Tên', 'Kiểu', 'Đích', 'Mô tả'],
        [
            ['output_1', 'pd.DataFrame', 'Module sau', '...'],
        ])
    
    doc.add_paragraph('Processing Logic (Pseudocode):', style='List Bullet')
    add_code_block(doc,
        'def process(inputs):\n'
        '    # Bước 1: ...\n'
        '    # Bước 2: ...\n'
        '    # Bước 3: ...\n'
        '    return output')
    
    doc.add_paragraph('Edge Cases & Error Handling:', style='List Bullet')
    add_table(doc,
        ['Tình huống', 'Cách xử lý'],
        [
            ['Input rỗng', 'Raise ValueError'],
            ['NaN trong tensor', 'Flag + skip'],
            ['OOM', 'Giảm batch size tự động'],
        ])
    
    doc.add_paragraph('Dependencies:', style='List Bullet')
    doc.add_paragraph('• Module: [Liệt kê]')
    doc.add_paragraph('• Library: [Liệt kê với version]')
    
    doc.add_paragraph('Estimated Complexity:', style='List Bullet')
    doc.add_paragraph('• Thời gian: [O(n), O(n²), ...]')
    doc.add_paragraph('• Bộ nhớ: [MB/GB cho input size N]')
    
    doc.add_page_break()
    
    # ===== PHẦN E =====
    add_heading_styled(doc, 'PHẦN E: DATA CONTRACTS', level=1)
    
    add_heading_styled(doc, 'E1. Data Classes', level=2)
    add_code_block(doc,
        '@dataclass(frozen=True)\n'
        'class SampleRecord:\n'
        '    sample_id: str          # Deterministic ID\n'
        '    class_name: str         # Ví dụ: "Apple___Apple_scab"\n'
        '    class_index: int        # 0-4\n'
        '    content_hash: str       # SHA-256\n'
        '    split: str              # "train" hoặc "test"')
    
    add_heading_styled(doc, 'E2. File Formats', level=2)
    add_table(doc,
        ['Artifact', 'Format', 'Schema', 'Immutable?'],
        [
            ['dataset_manifest', 'Parquet', 'Xem E1', '✅ YES'],
            ['stability_records', 'Parquet', 'Xem E3', '✅ YES'],
            ['statistical_results', 'Parquet', 'Xem E4', '✅ YES'],
        ])
    
    add_heading_styled(doc, 'E3. Naming Conventions', level=2)
    doc.add_paragraph('• File: snake_case.parquet')
    doc.add_paragraph('• Column: snake_case')
    doc.add_paragraph('• ID format: pv_<16-hex-chars>')
    
    add_heading_styled(doc, 'E4. Versioning', level=2)
    doc.add_paragraph('• Protocol version: 1.0 (frozen theo DR-001)')
    doc.add_paragraph('• Config hash: SHA-256 của resolved config')
    
    doc.add_page_break()
    
    # ===== PHẦN F =====
    add_heading_styled(doc, 'PHẦN F: BUSINESS RULES (QUAN TRỌNG NHẤT)', level=1)
    doc.add_paragraph(
        'Đây là phần QUAN TRỌNG NHẤT - dev/AI phải thuộc lòng. '
        'Vi phạm = PR bị REJECT ngay.'
    )
    
    add_heading_styled(doc, 'F1. Business Rules Table', level=2)
    add_table(doc,
        ['ID', 'Quy tắc', 'Module', 'Test ID'],
        [
            ['BR-001', 'Consistency filtering: chỉ evaluate nếu pred giống', 'JointEvaluator', 'T-FIL-01'],
            ['BR-002', 'Heatmap normalize về [0,1] với epsilon=1e-8', 'XAI', 'T-XAI-01'],
            ['BR-003', 'Transformation params từ YAML, không hardcode', 'Transform', 'T-TRF-01'],
            ['BR-007', 'Geometric transform phải có inverse alignment', 'JointEvaluator', 'T-ALN-01'],
            ['BR-008', 'sample_id chỉ dựa content hash, không dùng path', 'Data', 'T-DAT-01'],
            ['BR-010', 'Test set không dùng cho model selection', 'Training', 'T-TRN-01'],
            ['BR-013', 'Deterministic sign policy cho rotation', 'Transform', 'T-TRF-02'],
            ['BR-016', 'Statistical pairing bằng sample_id', 'Statistics', 'T-STA-01'],
            ['BR-017', 'Holm correction cho multiple comparisons', 'Statistics', 'T-STA-02'],
        ])
    
    add_heading_styled(doc, 'F2. Anti-patterns (KHÔNG ĐƯỢC)', level=2)
    doc.add_paragraph('❌ Hardcode đường dẫn tuyệt đối')
    doc.add_paragraph('❌ Dùng random.seed() không deterministic')
    doc.add_paragraph('❌ Commit weights/checkpoints lên Git')
    doc.add_paragraph('❌ Chỉnh số liệu thủ công trong paper')
    
    doc.add_page_break()
    
    # ===== PHẦN G =====
    add_heading_styled(doc, 'PHẦN G: SCIENTIFIC INVARIANTS', level=1)
    doc.add_paragraph(
        'Khác với business rules - đây là điều phải LUÔN ĐÚNG về mặt toán học/khoa học.'
    )
    add_table(doc,
        ['ID', 'Invariant', 'Cách verify'],
        [
            ['SI-01', 'Cùng seed + input → cùng output', 'Unit test'],
            ['SI-02', 'Heatmap ∈ [0, 1] sau normalize', 'Assertion'],
            ['SI-03', 'Rotate(θ) + Rotate(-θ) ≈ Identity (SSIM > 0.95)', 'Scientific test'],
            ['SI-04', 'SSIM(ảnh, ảnh) ≈ 1.0', 'Sanity test'],
            ['SI-05', 'p_adjusted ≥ p_raw (Holm monotonicity)', 'Unit test'],
            ['SI-06', 'Bootstrap CI chứa mean', 'Statistical test'],
            ['SI-07', 'Test set không bị chạm trong training', 'AST analysis'],
        ])
    
    doc.add_page_break()
    
    # ===== PHẦN H =====
    add_heading_styled(doc, 'PHẦN H: QUALITY GATES', level=1)
    add_table(doc,
        ['Gate', 'Stage', 'Criteria', 'Pass condition'],
        [
            ['G0', 'Environment setup', 'pip install -e . thành công', 'Exit code = 0'],
            ['G1', 'Data manifest', '5 classes, no cross-split dup', 'Validation pass'],
            ['G2', 'Transformations', '12 scenarios, deterministic', 'Tests pass'],
            ['G3', 'Model training', 'Acc ≥ 90%, no test leakage', 'Metrics OK'],
            ['G4', 'XAI', 'Heatmaps normalized', 'Tests pass'],
            ['G5', 'Joint Evaluator', 'BR-001 verified, pairing by ID', 'Scientific test'],
            ['G6', 'Statistics', 'Holm monotonicity', 'Reference match'],
            ['G7', 'Artifacts', '300 DPI, vector PDF', 'Visual check'],
            ['G8', 'Reproducibility', 'Clean rerun → same hashes', 'Hash match'],
            ['G9', 'Release', 'All tests pass, docs complete', 'CI green'],
        ])
    
    add_heading_styled(doc, 'H2. Definition of Done (DoD)', level=2)
    doc.add_paragraph('☐ Code có type hints đầy đủ')
    doc.add_paragraph('☐ Docstrings Google-style')
    doc.add_paragraph('☐ Unit tests coverage > 80%')
    doc.add_paragraph('☐ Scientific tests pass')
    doc.add_paragraph('☐ No hardcoded parameters')
    doc.add_paragraph('☐ PR reviewed & merged')
    
    doc.add_page_break()
    
    # ===== PHẦN I =====
    add_heading_styled(doc, 'PHẦN I: TESTING STRATEGY', level=1)
    
    add_heading_styled(doc, 'I1. Test Pyramid', level=2)
    add_code_block(doc,
        '        /\\\n'
        '       /  \\      Scientific Tests (5-10%)\n'
        '      / SI \\     → Verify scientific invariants\n'
        '     /------\\\n'
        '    /        /   Integration Tests (15-20%)\n'
        '   /  INT   /   → Verify module interactions\n'
        '  /--------/\n'
        ' /        /     Unit Tests (70-80%)\n'
        '/  UNIT  /     → Verify individual functions\n'
        '/--------/')
    
    add_heading_styled(doc, 'I2. Test Categories', level=2)
    doc.add_paragraph('Unit Tests (tests/unit/):', style='List Bullet')
    doc.add_paragraph('  - Test từng function/class độc lập, mock dependencies, coverage > 80%')
    doc.add_paragraph('Integration Tests (tests/integration/):', style='List Bullet')
    doc.add_paragraph('  - Test 2-3 modules làm việc cùng nhau, không mock')
    doc.add_paragraph('Scientific Tests (tests/scientific/) ⭐:', style='List Bullet')
    doc.add_paragraph('  - Test các SI, BẮT BUỘC cho research software')
    doc.add_paragraph('Reproducibility Tests (tests/reproducibility/):', style='List Bullet')
    doc.add_paragraph('  - Chạy 2 lần → so sánh artifact hashes (±1e-6)')
    
    doc.add_page_break()
    
    # ===== PHẦN J =====
    add_heading_styled(doc, 'PHẦN J: HANDOFF PROTOCOL', level=1)
    
    add_heading_styled(doc, 'J1. Handoff Checklist cho Developer', level=2)
    doc.add_paragraph('☐ Spec document hoàn chỉnh')
    doc.add_paragraph('☐ Reference implementations (nếu có)')
    doc.add_paragraph('☐ Test fixtures & golden files')
    doc.add_paragraph('☐ Decision Records (DR-xxx)')
    doc.add_paragraph('☐ Example prompts cho AI')
    
    add_heading_styled(doc, 'J2. Task Breakdown Template', level=2)
    add_code_block(doc,
        '### Task ID: TASK-[MODULE]-[N]\n'
        'Module: [Tên module]\n'
        'Priority: P0/P1/P2\n'
        'Estimated time: [X hours/days]\n'
        'Dependencies: [TASK-xxx]\n\n'
        'Mô tả: [Mô tả ngắn gọn]\n\n'
        'Acceptance Criteria:\n'
        '  - [ ] Criterion 1 (linked to BR-xxx)\n'
        '  - [ ] Criterion 2 (linked to SI-xxx)\n'
        '  - [ ] Tests pass\n'
        '  - [ ] Coverage > 80%\n\n'
        'Files to create/modify:\n'
        '  - src/.../file1.py (create)\n'
        '  - tests/.../test_file1.py (create)')
    
    add_heading_styled(doc, 'J3. AI Pair Programming Prompt Template', level=2)
    add_code_block(doc,
        '# PROMPT CHO AI\n\n'
        '## Vai trò\n'
        'Bạn là [Senior ML Engineer / Research Engineer]\n\n'
        '## Context\n'
        '- Dự án: [Tên]\n'
        '- Mục tiêu: [Mô tả]\n'
        '- Protocol version: [X.Y]\n\n'
        '## Nhiệm vụ\n'
        '[Mô tả chi tiết task]\n\n'
        '## Business Rules phải tuân thủ\n'
        '- BR-XXX: [Mô tả]\n\n'
        '## Data Contracts\n'
        '@dataclass class SomeRecord: ...\n\n'
        '## Yêu cầu kỹ thuật\n'
        '- Type hints đầy đủ\n'
        '- Docstrings Google-style\n'
        '- Unit tests với pytest\n\n'
        '## KHÔNG ĐƯỢC\n'
        '- Bịa số liệu\n'
        '- Thay đổi scientific protocol\n'
        '- Hardcode parameters')
    
    add_heading_styled(doc, 'J4. Review Checklist', level=2)
    doc.add_paragraph('☐ Code chạy được')
    doc.add_paragraph('☐ Tests pass')
    doc.add_paragraph('☐ Business rules được tuân thủ')
    doc.add_paragraph('☐ Scientific invariants được verify')
    doc.add_paragraph('☐ No overclaim trong comments/docs')
    doc.add_paragraph('☐ Traceability matrix được cập nhật')
    
    doc.add_page_break()
    
    # ===== HƯỚNG DẪN SỬ DỤNG =====
    add_heading_styled(doc, 'HƯỚNG DẪN SỬ DỤNG', level=1)
    
    add_heading_styled(doc, 'Cho Architect (người viết spec)', level=2)
    doc.add_paragraph('1. Điền đầy đủ 10 phần trước khi giao việc')
    doc.add_paragraph('2. Frozen protocol: Một khi spec đã chốt, KHÔNG thay đổi âm thầm')
    doc.add_paragraph('3. Decision Records: Mọi lựa chọn thiết kế đều phải có DR')
    doc.add_paragraph('4. Traceability: Mỗi RQ phải map được đến module → test → artifact')
    
    add_heading_styled(doc, 'Cho Developer (người implement)', level=2)
    doc.add_paragraph('1. Đọc spec 3 lần trước khi code')
    doc.add_paragraph('2. Thuộc lòng Business Rules (Phần F)')
    doc.add_paragraph('3. Chạy scientific tests sau mỗi module')
    doc.add_paragraph('4. Hỏi lại architect nếu spec không rõ')
    
    add_heading_styled(doc, 'Cho AI Pair Programming', level=2)
    doc.add_paragraph('1. Dùng prompt template (Phần J.3)')
    doc.add_paragraph('2. Cung cấp context đầy đủ: BR, SI, data contracts')
    doc.add_paragraph('3. Review code AI sinh - đừng tin mù quáng')
    doc.add_paragraph('4. Verify bằng tests')
    
    # ===== KẾT LUẬN =====
    doc.add_page_break()
    add_heading_styled(doc, 'KẾT LUẬN', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        'Template này không phải là giấy tờ hành chính - nó là "hợp đồng khoa học" '
        'giữa architect, developer, và reviewer.\n\n'
        '• Architect cam kết: "Đây là những gì chúng ta cần và tại sao"\n'
        '• Developer cam kết: "Tôi sẽ implement đúng như spec"\n'
        '• Reviewer cam kết: "Tôi sẽ verify rằng code đúng spec và spec đúng khoa học"\n\n'
        'Nếu thiếu template này, bạn sẽ có code chạy nhưng kết quả sai - '
        'điều tồi tệ nhất trong nghiên cứu.'
    )
    run.italic = True
    
    # ===== SAVE =====
    output_path = 'Template_DacTa_ResearchSoftware.docx'
    doc.save(output_path)
    print(f'✅ File đã được tạo: {output_path}')
    print(f'📂 Mở file bằng Microsoft Word hoặc LibreOffice Writer')


if __name__ == '__main__':
    build_document()