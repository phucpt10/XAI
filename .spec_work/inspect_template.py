from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(r"D:\ResearchCode\XAI")
DOCX = ROOT / "Template_DacTa_ResearchSoftware.docx"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


doc = Document(DOCX)
report = {
    "reference": str(DOCX),
    "sha256": sha256(DOCX),
    "paragraphs": len(doc.paragraphs),
    "tables": len(doc.tables),
    "sections": len(doc.sections),
    "inline_shapes": len(doc.inline_shapes),
    "styles": {},
    "headers": [[p.text for p in s.header.paragraphs] for s in doc.sections],
    "footers": [[p.text for p in s.footer.paragraphs] for s in doc.sections],
}

for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
    if name not in doc.styles:
        continue
    style = doc.styles[name]
    pf = style.paragraph_format
    font = style.font
    report["styles"][name] = {
        "font": font.name,
        "size_pt": font.size.pt if font.size else None,
        "bold": font.bold,
        "italic": font.italic,
        "color": str(font.color.rgb) if font.color and font.color.rgb else None,
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "line_spacing": pf.line_spacing,
        "left_indent_in": pf.left_indent.inches if pf.left_indent else None,
        "first_line_indent_in": pf.first_line_indent.inches if pf.first_line_indent else None,
        "keep_with_next": pf.keep_with_next,
    }

with zipfile.ZipFile(DOCX) as archive:
    report["package_parts"] = [
        {"path": i.filename, "size": len(archive.read(i.filename)), "sha256": hashlib.sha256(archive.read(i.filename)).hexdigest()}
        for i in archive.infolist()
    ]
    report["app_xml"] = archive.read("docProps/app.xml").decode("utf-8", errors="replace")

print(json.dumps(report, ensure_ascii=False, indent=2, default=str))
