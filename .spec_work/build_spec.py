from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\ResearchCode\XAI")
REFERENCE = ROOT / "Template_DacTa_ResearchSoftware.docx"
OUTPUT = ROOT / "PlantXAI-Stability_Research_Software_Specification_v0.9_Draft.docx"
REFERENCE_SHA256 = "44941816fee5f0c6e4d76a0157b3d11fd1bd0e4b767f6ff779fbefa031944b61"

NAVY = "1F4E79"
BLUE = "4F81BD"
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "F2F2F2"
LIGHT_AMBER = "FFF2CC"
LIGHT_RED = "FCE4D6"
LIGHT_GREEN = "E2F0D9"
WHITE = "FFFFFF"
TEXT = "222222"
TOTAL_DXA = 8640  # 6.00 in usable width from retained template


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def set_run_font(run, name: str = "Calibri", size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)


def shade(container, fill: str) -> None:
    pr = container._tc.get_or_add_tcPr() if hasattr(container, "_tc") else container._p.get_or_add_pPr()
    old = pr.find(qn("w:shd"))
    if old is not None:
        pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def set_cell_margins(cell, top: int = 100, start: int = 100, bottom: int = 100, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != TOTAL_DXA:
        raise ValueError(f"Table widths must sum to {TOTAL_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOTAL_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_keep(paragraph, keep_next: bool = False, keep_lines: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    if keep_lines:
        keep = OxmlElement("w:keepLines")
        p_pr.append(keep)


def add_text(doc: Document, text: str = "", *, bold_prefix: str | None = None, italic: bool = False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        set_run_font(r1)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
        r2.italic = italic
    else:
        run = p.add_run(text)
        set_run_font(run)
        run.italic = italic
    return p


def add_bullets(doc: Document, items: Iterable[str], level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    for item in items:
        p = doc.add_paragraph(style=style if style in doc.styles else "List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            set_run_font(run)
        if not p.runs:
            run = p.add_run(item)
            set_run_font(run)
        else:
            p.runs[0].text = item
        set_paragraph_keep(p, keep_lines=True)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item) if not p.runs else p.runs[0]
        run.text = item
        set_run_font(run)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    return p


def major_section(doc: Document, text: str) -> None:
    doc.add_page_break()
    heading(doc, text, 1)


def callout(doc: Document, label: str, text: str, fill: str = LIGHT_AMBER) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    shade(p, fill)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(NAVY)
    set_run_font(label_run)
    text_run = p.add_run(text)
    set_run_font(text_run)


def code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.30)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    shade(p, LIGHT_GREY)
    run = p.add_run(text)
    set_run_font(run, "Consolas", 9)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[object]], widths: Sequence[int], font_size: float = 9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(header)
        shade(cell, NAVY)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=font_size)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
    repeat_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = str(value)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, size=font_size)
                    run.font.color.rgb = RGBColor.from_string(TEXT)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def module_spec(
    doc: Document,
    identifier: str,
    name: str,
    purpose: str,
    inputs: Sequence[str],
    outputs: Sequence[str],
    logic: Sequence[str],
    failures: Sequence[str],
    acceptance: Sequence[str],
) -> None:
    heading(doc, f"{identifier}. {name}", 2)
    add_text(doc, f"Purpose: {purpose}")
    heading(doc, "Inputs", 3)
    add_bullets(doc, inputs)
    heading(doc, "Outputs", 3)
    add_bullets(doc, outputs)
    heading(doc, "Processing Logic", 3)
    add_numbered(doc, logic)
    heading(doc, "Failure and Exclusion Conditions", 3)
    add_bullets(doc, failures)
    heading(doc, "Acceptance Criteria", 3)
    add_bullets(doc, acceptance)


def contract_table(doc: Document, rows: Sequence[Sequence[object]]) -> None:
    add_table(doc, ["Field", "Type", "Required", "Description / Constraint"], rows, [1750, 1450, 900, 4540], font_size=8.0)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 3"].font.size = Pt(11)


def add_title_page(doc: Document) -> None:
    title = doc.add_heading("PLANTXAI-STABILITY\nRESEARCH SOFTWARE DESIGN & SPECIFICATION", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    for run in title.runs:
        set_run_font(run, size=24)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("17365D")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    run = subtitle.add_run("Developer Edition\nResearch Software Architecture, Scientific Protocol, Data Contracts, QA, and Handoff")
    set_run_font(run, size=12)
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Document version", "0.9 Draft"],
            ["Scientific protocol", "v0.9 - DRAFT; not frozen"],
            ["Readiness", "G0A PASS; G0B BLOCKED"],
            ["Prepared", "22 July 2026"],
            ["Audience", "Research lead, software architect, developers, reviewers, and AI pair-programming agents"],
            ["Authority", "Human research lead retains final approval for all scientific choices"],
        ],
        [2200, 6440],
        font_size=9.5,
    )
    callout(
        doc,
        "DRAFT CONTROL",
        "This specification is implementation-ready only for software scaffolding and validation infrastructure. Official dataset execution, training, CAM generation, statistical analysis, and research claims remain blocked until G0B passes and protocol v1 is approved and frozen.",
        LIGHT_RED,
    )
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    heading(doc, "DESIGN PHILOSOPHY", 1)
    add_text(
        doc,
        "PlantXAI-Stability is a reproducible research experimentation framework, not a real-time crop-disease diagnostic product. Its primary quality attributes are scientific validity, traceability, deterministic identity, explicit exclusion handling, and reproducible artifact generation.",
    )
    add_table(
        doc,
        ["Concern", "Conventional Software", "Research Software Requirement"],
        [
            ["Goal", "Correct execution and scale", "Correct execution plus reproducible and auditable scientific evidence"],
            ["Primary users", "Operators and developers", "Researchers, reviewers, developers, and the scientific community"],
            ["Change model", "Iterative feature delivery", "Versioned protocol; scientific changes require decisions and new run identities"],
            ["Quality", "Bug reduction", "Traceable data lineage, testable invariants, and claim boundaries"],
            ["Output", "Deployed feature", "Per-sample records, statistics, publication artifacts, and provenance bundle"],
        ],
        [1550, 2750, 4340],
    )
    callout(doc, "CORE PRINCIPLE", "Freeze scientific intent before official execution. Code that runs but produces invalid evidence is a failed research artifact.", LIGHT_BLUE)

    heading(doc, "Document Status Vocabulary", 2)
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["APPROVED", "A human-approved rule that implementations must enforce."],
            ["PROVISIONAL", "A candidate value requiring pilot evidence and approval before protocol freeze."],
            ["PENDING", "An unresolved scientific or operational choice."],
            ["BLOCKED", "Execution is prohibited until the stated evidence and approval exist."],
        ],
        [1800, 6840],
    )

    heading(doc, "Specification Structure", 2)
    add_table(
        doc,
        ["Part", "Name", "Purpose"],
        [
            ["A", "Research Context", "Why the study exists and what questions it answers"],
            ["B", "Scope and Boundaries", "Allowed work, excluded work, claims, and limitations"],
            ["C", "System Architecture", "Layers, packages, pipelines, dependencies, and technologies"],
            ["D", "Module Specifications", "Responsibilities, inputs, outputs, algorithms, and failures"],
            ["E", "Data Contracts", "Typed records, keys, immutability, storage, and schemas"],
            ["F", "Business Rules", "Non-negotiable implementation policies"],
            ["G", "Scientific Invariants", "Properties that must always hold scientifically"],
            ["H", "Quality Gates", "Evidence required to proceed between stages"],
            ["I", "Testing Strategy", "Unit, integration, scientific, and reproducibility tests"],
            ["J", "Handoff Protocol", "Tasks, acceptance criteria, change control, and traceability"],
        ],
        [850, 2300, 5490],
    )


def build_part_a(doc: Document) -> None:
    major_section(doc, "PART A - RESEARCH CONTEXT")
    add_text(doc, "Purpose: define why the research is conducted, what evidence is required, and which claims the software may support.", italic=True)

    heading(doc, "A1. Research Title", 2)
    add_text(doc, "PlantXAI-Stability: Joint Evaluation of Prediction Robustness and Explanation Stability for Plant-Disease Classification under Controlled Image Transformations")

    heading(doc, "A2. Problem Statement", 2)
    add_text(doc, "Image classifiers can preserve their predicted class after rotation, brightness variation, blur, or noise while changing the image regions used to support that decision. Conventional performance reporting focuses on accuracy and may therefore miss unstable visual evidence. The study requires a controlled, paired, and reproducible protocol that separates output robustness from explanation stability.")
    callout(doc, "SCIENTIFIC GAP", "A stable prediction is not evidence of a stable explanation. The model may return the same class while relying on different visual regions after a small, semantically preserving transformation.", LIGHT_BLUE)

    heading(doc, "A3. Research Questions", 2)
    add_table(
        doc,
        ["ID", "Research Question", "Primary Modules", "Required Evidence"],
        [
            ["RQ1", "How robust are ResNet50 and EfficientNet-B0 predictions under controlled image transformations?", "InferenceEngine; PredictionRobustness; JointEvaluator", "Per-pair class consistency, confidence change, severity trends, and cluster-aware uncertainty"],
            ["RQ2", "How do Grad-CAM, Grad-CAM++, and Score-CAM differ in explanation stability?", "XAIGenerator; HeatmapAligner; ExplanationMetrics; StatisticalAnalysis", "Aligned heatmap pairs, SSIM, Pearson, cosine, paired tests, effect sizes, and adjusted p-values"],
            ["RQ3", "What relationship exists between prediction robustness and explanation stability?", "JointEvaluator; RelationshipAnalyzer; StatisticalAnalysis", "Joint per-sample records and a predeclared analysis of confidence change versus stability in the prediction-consistent population"],
        ],
        [700, 3150, 2100, 2690],
        font_size=8.5,
    )

    heading(doc, "A4. Expected Contributions", 2)
    add_bullets(doc, [
        "A controlled protocol that evaluates prediction robustness and explanation stability as distinct but linked properties.",
        "A paired benchmark of three CAM methods across two CNN backbones and twelve controlled transformation scenarios.",
        "A transparent eligibility and exclusion model that prevents class-switching and invalid heatmaps from contaminating stability estimates.",
        "Leaf-aware uncertainty estimates, paired hypothesis testing, effect sizes, and family-wise multiple-testing control.",
        "A reproducible research-software artifact with immutable manifests, hashes, frozen splits, configuration lineage, per-sample records, and regenerable publication tables and figures.",
    ])
    add_text(doc, "The study does not propose a new XAI algorithm. Its contribution is the experimental protocol, software architecture, comparative evidence, and reproducibility system.")

    heading(doc, "A5. Study Population and Scientific Factors", 2)
    add_table(
        doc,
        ["Factor", "Planned Levels", "Status"],
        [
            ["Dataset", "PlantVillage color images; immutable revision required", "PENDING audit and immutable revision"],
            ["Classes", "Five tomato classes including healthy control", "PROVISIONAL until class audit approval"],
            ["Backbones", "ResNet50; EfficientNet-B0", "PROVISIONAL until checkpoints are selected"],
            ["XAI methods", "Grad-CAM; Grad-CAM++; Score-CAM", "PLANNED"],
            ["Transformations", "Rotation; brightness; Gaussian noise; Gaussian blur", "PROVISIONAL pending semantic pilot"],
            ["Severities", "Mild; moderate; severe", "PROVISIONAL values in Appendix A"],
            ["Metrics", "SSIM; Pearson correlation; cosine similarity", "PLANNED; exact implementations to be frozen"],
        ],
        [1800, 4700, 2140],
    )

    heading(doc, "A6. Audience and Intended Use", 2)
    add_bullets(doc, [
        "Research leads who approve the scientific protocol and claims.",
        "Research software engineers implementing the pipeline and quality gates.",
        "Reviewers auditing data lineage, exclusions, statistical pairing, and reproducibility.",
        "AI pair-programming systems operating under explicit task scope, data contracts, and non-negotiable scientific rules.",
    ])
    add_text(doc, "Target publication venue and ranking are PENDING and must be recorded without changing the approved experiment after official execution begins.")

    heading(doc, "A7. Claim Boundary", 2)
    add_table(
        doc,
        ["Allowed Claim Form", "Prohibited Overclaim"],
        [
            ["On the approved five-class PlantVillage population, model A retained predictions more often than model B under transformation X.", "Model A is robust for real-world agricultural diagnosis."],
            ["CAM A produced higher paired stability than CAM B for declared scenarios after correction and effect-size reporting.", "CAM A is universally better or more faithful than CAM B."],
            ["Stable predictions were not always accompanied by stable heatmaps in the eligible population.", "The study proves a causal mechanism for model reasoning."],
        ],
        [4320, 4320],
    )


def build_part_b(doc: Document) -> None:
    major_section(doc, "PART B - SCOPE AND BOUNDARIES")

    heading(doc, "B1. In Scope", 2)
    add_bullets(doc, [
        "Auditing an immutable PlantVillage dataset revision and establishing canonical RGB identities.",
        "Selecting five tomato classes using audit evidence rather than test performance.",
        "Creating leaf-grouped, class-stratified frozen train/validation/test splits while preserving the official test policy.",
        "Training and validation-based checkpoint selection for ResNet50 and EfficientNet-B0.",
        "Applying twelve deterministic transformation scenarios and recording exact per-sample parameters.",
        "Generating Grad-CAM, Grad-CAM++, and Score-CAM explanations with an explicit common target class.",
        "Inverse-aligning geometric heatmaps and calculating SSIM, Pearson correlation, and cosine similarity.",
        "Conducting leaf-aware bootstrap uncertainty, paired Wilcoxon tests, rank-biserial effect sizes, and Holm correction.",
        "Generating machine-readable records, paper tables, figures, audit reports, and reproducibility manifests.",
    ])

    heading(doc, "B2. Out of Scope", 2)
    add_bullets(doc, [
        "Real-time disease diagnosis, deployment APIs, mobile applications, or clinical/agronomic decision support.",
        "Claims of field-domain validity, cross-crop generalization, cross-dataset transfer, or production readiness.",
        "Development of a new XAI algorithm.",
        "A definitive assessment of explanation faithfulness, causal validity, localization accuracy, or human usefulness.",
        "Using the official test set for class selection, severity tuning, target-layer selection, hyperparameter tuning, or checkpoint selection.",
        "Manual editing of computed results in the manuscript, tables, or figures.",
    ])

    heading(doc, "B3. Assumptions and Preconditions", 2)
    add_table(
        doc,
        ["ID", "Assumption / Precondition", "Failure Response"],
        [
            ["AS-01", "An immutable dataset revision can be identified and retrieved.", "Block G1; do not build an official manifest."],
            ["AS-02", "A trustworthy leaf grouping identifier exists or can be established from documented source evidence.", "Block split freeze and leaf-cluster inference; never invent leaf IDs."],
            ["AS-03", "Transformations selected for the protocol preserve disease semantics at approved severities.", "Revise or reject severity based only on train/validation pilot evidence."],
            ["AS-04", "Approved model target layers expose spatial activations and required gradients.", "Block G4 for the affected model-method combination."],
            ["AS-05", "The runtime can record exact package, hardware, seed, Git, configuration, and artifact provenance.", "Official run is prohibited."],
        ],
        [900, 5000, 2740],
        font_size=8.8,
    )

    heading(doc, "B4. Approved and Provisional Scientific Choices", 2)
    add_table(
        doc,
        ["Choice", "Current State", "Freeze Evidence"],
        [
            ["Protocol v0.9", "APPROVED as draft only", "Human approval of all seven G0B blockers and canonical protocol hash"],
            ["Five tomato classes", "PROVISIONAL", "Dataset audit, leaf/image counts, duplication analysis, and class-selection decision record"],
            ["Official test preservation", "PLANNED", "Verified dataset split policy and non-overlap audit"],
            ["Transformation severities", "PROVISIONAL", "Train/validation pilot plus semantic review"],
            ["Training hyperparameters", "PROVISIONAL", "Predeclared tuning policy using train/validation only"],
            ["Target layers", "PROVISIONAL", "Runtime activation-shape, gradient, hook-cleanup, and smoke-heatmap evidence"],
            ["Statistical families", "PENDING", "Predeclared comparison-family decision record"],
        ],
        [2250, 1850, 4540],
    )

    heading(doc, "B5. Known Limitations", 2)
    add_bullets(doc, [
        "PlantVillage backgrounds and acquisition conditions can produce shortcuts that remain after same-crop selection.",
        "Explanation stability is observable only for prediction-consistent pairs, producing a selected population for RQ3.",
        "CAM stability does not establish explanation correctness or faithfulness.",
        "GPU operations may be numerically reproducible only within a declared tolerance rather than bit-exact.",
        "The effective sample size may be smaller than the image count when multiple images originate from one leaf.",
        "Severe transformations may alter disease evidence; pilot approval is therefore a scientific gate, not an implementation detail.",
    ])

    heading(doc, "B6. Execution Authorization", 2)
    callout(doc, "BLOCKED", "No official dataset execution, model training, CAM generation, statistical testing, or research-result generation is authorized while G0B remains blocked. Software bootstrap, schema validation, test fixtures, and non-scientific smoke tests may proceed within approved increments.", LIGHT_RED)


def build_part_c(doc: Document) -> None:
    major_section(doc, "PART C - SYSTEM ARCHITECTURE")

    heading(doc, "C1. Architectural Style", 2)
    add_text(doc, "The system uses a layered, modular, configuration-driven architecture. Scientific domain logic resides in typed package modules, not in notebooks or CLI entry points. Application services coordinate work; infrastructure adapters provide PyTorch, filesystem, Parquet, plotting, and environment access.")
    code_block(doc, "CLI / Experiment Scripts\n        |\nApplication Orchestration: Trainer, JointEvaluator\n        |\nScientific Domain Services: Prediction, XAI, Alignment, Metrics, Statistics\n        |\nTyped Interfaces and Immutable Data Contracts\n        |\nInfrastructure: PyTorch, Filesystem, Parquet, Plotting\n\nCross-cutting: Configuration | Provenance | QA | Atomic Resume")

    heading(doc, "C2. Scientific Data Flow", 2)
    code_block(doc, "Frozen manifest and splits\n  -> canonical RGB input\n  -> original prediction\n  -> deterministic transformed input\n  -> transformed prediction\n  -> prediction record and consistency gate\n       inconsistent -> retain for RQ1; document exclusion from stability\n       consistent   -> paired CAMs with original predicted target\n  -> inverse alignment and valid-overlap masking\n  -> heatmap normalization and quality validation\n  -> SSIM / Pearson / cosine per-sample records\n  -> joint population reconciliation\n  -> leaf-cluster bootstrap + paired Wilcoxon + effect size + Holm\n  -> validated tables, figures, reports, and reproducibility bundle")

    heading(doc, "C3. Package Map and Dependency Boundaries", 2)
    add_table(
        doc,
        ["Package", "Responsibility", "Must Not Contain"],
        [
            ["config", "Schema, validation, merge, resolution, and immutable configuration hash", "Dataset I/O, inference, plotting, or silent fallback"],
            ["data", "Dataset receipt, canonical identity, manifest, split, dataset adapter, and DataLoader", "Training, XAI, or paper statistics"],
            ["transformations", "Deterministic transformations, parameter records, inverse metadata, and valid masks", "Model or statistical logic"],
            ["models", "Model wrappers, training, checkpoint selection, identity, and inference", "Paper plotting or dataset splitting"],
            ["xai", "CAM adapters, hooks, target-layer resolution, heatmap generation, and basic quality metadata", "Split creation or statistical tests"],
            ["evaluation", "Prediction comparison, consistency policy, alignment, metric orchestration, and joint records", "Hard-coded scientific parameters"],
            ["statistics", "Cluster bootstrap, paired Wilcoxon, effect size, Holm, and relationship analysis", "Image transforms or model inference"],
            ["visualization", "Tables, figures, qualitative selections, and reports from validated records", "Recomputation of scientific metrics"],
            ["provenance", "Run identity, environment, Git, hashes, lineage, status, and artifact index", "Scientific decision logic"],
            ["cli", "Thin entry points that invoke application services", "Business or scientific rules"],
            ["utils", "Logging, device, seed, time, and common I/O primitives", "Protocol decisions"],
        ],
        [1550, 4070, 3020],
        font_size=8.0,
    )

    heading(doc, "C4. Central Orchestration", 2)
    add_text(doc, "JointEvaluator is the central application service. It owns call order, consistency-policy enforcement, common target propagation, transactional result handling, and exclusions. It must not implement transformations, CAM algorithms, similarity metrics, hypothesis tests, or plotting.")
    code_block(doc, "JointEvaluator\n  +- ConfigLoader / ProtocolValidator\n  +- DatasetManifestBuilder / SplitManager / DatasetAdapter\n  +- TransformationPipeline\n  +- ModelRegistry / Trainer / InferenceEngine\n  +- XAIGenerator\n  +- HeatmapAligner\n  +- ExplanationMetrics\n  +- RelationshipAnalyzer / StatisticalAnalysis\n  +- ResultStore / ArtifactGenerator\n  +- RunManifestWriter")

    heading(doc, "C5. High-Level Execution Gates", 2)
    add_table(
        doc,
        ["Gate", "Stage", "Required Transition"],
        [
            ["G0", "Configuration and protocol", "Resolve typed configuration; confirm authorization and hashes"],
            ["G1", "Dataset", "Validate receipt, classes, canonical manifest, grouping, and frozen splits"],
            ["G2", "Models", "Load validation-selected checkpoints with complete evidence"],
            ["G3", "Transformations", "Confirm twelve deterministic, semantically approved scenarios"],
            ["G4", "XAI", "Validate three CAM adapters and target layers"],
            ["G5", "Alignment and metrics", "Pass synthetic alignment and metric reference tests"],
            ["G6", "Joint records", "Reconcile keys, populations, exclusions, and atomic partitions"],
            ["G7", "Statistics", "Validate pairing, cluster units, tests, effect sizes, and correction families"],
            ["G8", "Artifacts", "Regenerate tables and figures from validated records"],
            ["G9", "Reproducibility", "Clean rerun and artifact-hash/tolerance comparison"],
        ],
        [900, 2400, 5340],
    )

    heading(doc, "C6. Technology Stack", 2)
    add_table(
        doc,
        ["Layer", "Candidate Technology", "Version Policy", "Rationale / Status"],
        [
            ["Language", "Python", "Pinned supported interpreter; repository currently targets Python 3.11.2", "Typed research implementation; exact environment recorded"],
            ["Deep learning", "PyTorch / torchvision", "Pinned before official run", "Model training, inference, and activation/gradient access"],
            ["Dataset", "Approved dataset client plus Pillow/array stack", "Pinned before data receipt", "Immutable revision retrieval and canonical RGB conversion"],
            ["XAI", "Validated in-house adapters and/or pinned CAM library", "PENDING choice", "Uniform target-class and target-layer contract"],
            ["Tabular storage", "Parquet and JSON", "Schema-versioned", "Per-sample evidence and run metadata"],
            ["Statistics", "SciPy/statsmodels plus validated project code", "Pinned before analysis", "Wilcoxon, effect size, correction, and reference tests"],
            ["Testing", "pytest, Ruff, MyPy", "Repository-pinned", "Static, unit, integration, scientific, and reproducibility QA"],
        ],
        [1450, 2100, 2050, 3040],
        font_size=8.2,
    )

    heading(doc, "C7. Repository Structure", 2)
    code_block(doc, "project/\n+- configs/                 # Draft/frozen protocol and versioned component settings\n+- src/plantxai_stability/ # Package implementation\n|  +- config/ data/ transformations/ models/ xai/ evaluation/\n|  +- statistics/ visualization/ provenance/ cli/ utils/\n+- tests/\n|  +- unit/ integration/ scientific/ reproducibility/ fixtures/\n+- scripts/                 # Thin operational utilities\n+- docs/                    # Specifications and decision records\n+- data/                    # Ignored raw/interim data; approved manifests/splits by policy\n+- outputs/runs/<run_id>/   # Immutable run-scoped artifacts\n+- pyproject.toml           # Tooling and dependencies")

    heading(doc, "C8. Runtime Storage and Resume", 2)
    add_text(doc, "Large evaluation work is partitioned by model, scenario, and XAI method. Each partition uses atomic write-then-rename semantics, a unique composite key, checksum, completion marker, and configuration identity. Resume executes only missing or invalid partitions and never appends incompatible data to an existing run.")
    code_block(doc, "outputs/runs/<run_id>/\n+- resolved_config.yaml\n+- run_manifest.json\n+- predictions/\n+- heatmaps/\n+- metrics/\n+- statistics/\n+- tables/\n+- figures/\n+- reports/\n+- logs/\n+- status/\n+- artifact_index.json")


def build_part_d(doc: Document) -> None:
    major_section(doc, "PART D - MODULE SPECIFICATIONS")
    add_text(doc, "Each module has one primary responsibility, typed inputs and outputs, explicit failure behavior, and acceptance criteria linked to Parts F-I. Scientific parameters must enter through validated configuration rather than module defaults.", italic=True)

    module_spec(
        doc,
        "D1",
        "ConfigLoader and ProtocolValidator",
        "Resolve all versioned configuration into an immutable, typed, hashable scientific execution contract.",
        ["Protocol source path and component YAML paths", "Environment-independent command arguments", "Schema and cross-field validation rules"],
        ["ResolvedConfig object", "Canonical resolved_config.yaml", "Configuration bundle SHA-256", "G0 readiness report"],
        ["Load only explicitly named protocol sources; do not infer a future protocol version.", "Reject unknown keys, invalid types, illegal enum values, absolute scientific artifact paths, and inconsistent cross-field choices.", "Canonicalize the resolved configuration, compute its hash, and make the runtime object immutable.", "Evaluate G0A/G0B without modifying protocol files."],
        ["Unknown key or schema error", "Frozen protocol with unresolved blockers", "Configuration hash mismatch", "Missing required scientific component"],
        ["Same inputs produce the same canonical hash.", "Draft protocol may represent pending evidence; frozen protocol fails closed.", "No module receives unvalidated raw dictionaries."],
    )

    module_spec(
        doc,
        "D2",
        "DatasetReceipt and DatasetManifestBuilder",
        "Establish an auditable dataset source and deterministic identity for every canonical image.",
        ["Approved dataset repository/configuration/revision", "Dataset feature and label schemas", "Canonical RGB conversion policy", "Approved class-selection record"],
        ["dataset_receipt.json", "dataset_schema.json", "class_mapping.json", "dataset_manifest.parquet", "Duplicate and integrity reports", "Manifest SHA-256"],
        ["Verify the immutable upstream and operational revision.", "Decode labels using the source schema and convert images to canonical RGB uint8 HWC.", "Hash canonical pixels and construct stable sample_id values independent of row order and absolute paths.", "Audit corruption, exact duplicates, conflicting labels, class availability, and leaf grouping.", "Sort deterministically and finalize the manifest only after approval."],
        ["Mutable or unidentified revision", "Corrupted input", "Duplicate or conflicting identity", "Missing approved class", "Missing/untrustworthy leaf_id evidence"],
        ["Manifest rerun on the same revision has identical membership and hash.", "Every row has sample_id, leaf_id, class, source split, path, dimensions, and canonical hash.", "No silent deletion; every rejected source record has a reason."],
    )

    module_spec(
        doc,
        "D3",
        "SplitManager, DatasetAdapter, and DataLoaderFactory",
        "Create leakage-safe frozen splits and preserve identity through model-ready batching.",
        ["Approved manifest", "Official source split policy", "Grouping and stratification configuration", "Preprocessing configuration"],
        ["train/validation/test split files and hashes", "Split leakage report", "Typed SampleRecord batches", "Data quality report"],
        ["Preserve the official test population and split only approved training data when required.", "Group by leaf_id and stratify by class without cross-split group overlap.", "Freeze membership and hashes; official adapters only read frozen split files.", "Return pixel_tensor in [0,1] plus model_tensor, label, sample_id, leaf_id, split, and source metadata.", "Seed workers; allow train shuffle only; disable shuffle and drop_last for validation/test inference."],
        ["Sample or leaf overlap", "Missing manifest row", "Class-index mismatch", "Non-RGB/corrupt image", "Attempt to recreate an official frozen split implicitly"],
        ["Zero sample_id and leaf_id overlap across splits.", "Repeated unshuffled iteration yields the same identity sequence.", "Transformations operate before model normalization."],
    )

    module_spec(
        doc,
        "D4",
        "TransformationPipeline",
        "Apply controlled, deterministic perturbations and return complete inverse/provenance metadata.",
        ["Canonical pixel_tensor in [0,1]", "sample_id", "ScenarioConfig", "Global seed"],
        ["Transformed pixel_tensor", "TransformationRecord", "Forward/inverse metadata", "Valid-overlap mask metadata when geometric"],
        ["Resolve one of twelve approved scenarios.", "Derive a per-sample seed from global seed, sample_id, and scenario identity using a stable algorithm.", "Apply rotation, brightness, Gaussian noise, or Gaussian blur before model normalization.", "Record exact direction, value, interpolation, fill/border policy, and seed.", "Clip valid pixel ranges and return data sufficient to reproduce the exact input."],
        ["Unknown scenario", "Non-finite transformed pixels", "Missing inverse metadata for rotation", "Semantic-pilot status not approved for official run"],
        ["Same input and scenario reproduce identical transformed pixels and metadata.", "Both models receive the same transformed input.", "All twelve scenarios pass synthetic and semantic-pilot gates before official use."],
    )

    module_spec(
        doc,
        "D5",
        "ModelRegistry, Trainer, and InferenceEngine",
        "Provide uniform training, checkpoint evidence, identity, and prediction records for both CNN backbones.",
        ["Typed model/training configuration", "Frozen train/validation splits", "Checkpoint reference and expected hash", "Model tensor batch"],
        ["Selected checkpoint and evidence", "Training history", "Logits/probabilities/confidence", "PredictionRecord", "Model and checkpoint identity"],
        ["Instantiate an approved pretrained backbone and replace its classifier for five classes.", "Run smoke and tiny-batch overfit checks before full training.", "Train only on train; select checkpoint using the predeclared validation metric.", "Verify checkpoint hash before inference and set evaluation/no-gradient mode.", "Emit per-sample predictions keyed by sample_id; never join by batch position."],
        ["Test data accessed during selection", "Checkpoint hash mismatch", "Class mapping mismatch", "Non-finite logits", "Unsupported target-layer mapping"],
        ["Selection evidence includes validation metric, seed, training-config hash, Git SHA, and artifact hash.", "Original and transformed inference use the identical approved checkpoint.", "Official test inference is deterministic within declared tolerance."],
    )

    module_spec(
        doc,
        "D6",
        "XAIGenerator",
        "Generate comparable Grad-CAM, Grad-CAM++, and Score-CAM heatmaps through one explicit interface.",
        ["Model wrapper and approved checkpoint", "Input model tensor", "Caller-supplied target_class", "Runtime-validated target_layer", "XAI method configuration"],
        ["Two-dimensional heatmap", "CAMRecord", "Activation/gradient quality metadata", "Generation timing and checksum"],
        ["Resolve and validate the configured target layer at runtime.", "Attach method-specific hooks, generate the heatmap, and remove hooks in a guaranteed cleanup block.", "Resize heatmap to the declared input resolution without changing the scientific target.", "For Score-CAM, use deterministic chunking/cache that changes execution cost but not method semantics.", "Return raw heatmap plus quality metadata; do not silently normalize invalid output."],
        ["Invalid target layer", "Missing gradients/activations", "NaN/Inf output", "Constant/near-constant output", "Hook cleanup failure", "Out-of-memory CAM failure"],
        ["Target class is never inferred internally.", "All methods use the same class for a paired comparison.", "Repeated hook use shows no accumulation or cross-sample leakage."],
    )

    module_spec(
        doc,
        "D7",
        "HeatmapAligner",
        "Place transformed heatmaps in the original coordinate frame and define the valid comparison region.",
        ["Raw transformed heatmap", "Original heatmap shape", "TransformationRecord inverse metadata", "Mask policy"],
        ["Aligned heatmap", "Valid-overlap mask", "Alignment audit metadata"],
        ["Use identity alignment for brightness, noise, and blur.", "For rotation, apply the exact inverse angle and spatial convention from the transformation record.", "Resize to the original heatmap shape using the frozen interpolation policy.", "Transform a support mask and exclude padding/border pixels from comparison.", "Validate that the mask is non-empty and the aligned output is finite."],
        ["Missing/invalid inverse metadata", "Angle-sign or convention mismatch", "Empty valid mask", "Shape mismatch", "Non-finite output"],
        ["Synthetic hotspot tests show higher similarity after correct inverse alignment.", "No alignment parameter is estimated from the heatmap itself.", "Image and heatmap transformations use compatible spatial conventions."],
    )

    module_spec(
        doc,
        "D8",
        "HeatmapQualityValidator and ExplanationMetrics",
        "Normalize valid heatmaps and compute complementary per-sample stability metrics without fabricating values.",
        ["Original and aligned transformed heatmaps", "Valid-overlap mask", "Metric configuration"],
        ["Normalized heatmaps or documented exclusion", "SSIM, Pearson, and cosine values", "StabilityRecord"],
        ["Validate two-dimensional shape, common resolution, finite values, mask support, and variance/norm conditions.", "Min-max normalize each valid heatmap to [0,1] using epsilon 1e-8.", "Reject constant or near-constant maps before metrics.", "Compute SSIM with data_range=1.0 and Pearson/cosine on valid-mask pixels.", "Persist per-sample metrics and quality provenance before any aggregation."],
        ["constant_heatmap", "non_finite_heatmap", "empty_valid_mask", "zero_variance_pearson", "zero_norm_cosine", "shape_mismatch"],
        ["Identical valid maps produce metrics approximately one.", "Invalid maps never receive zero or perfect similarity by fallback.", "No metric rows disappear without a matching exclusion record."],
    )

    module_spec(
        doc,
        "D9",
        "JointEvaluator and PredictionRobustness",
        "Coordinate paired inference, enforce the consistency gate, propagate the common target, and reconcile populations.",
        ["Frozen test SampleRecords", "Approved model/checkpoint identities", "Twelve scenarios", "Three XAI methods", "Validated services and ResultStore"],
        ["PredictionRecord for every inferable pair", "Consistency mask", "JointRecord", "ExclusionRecord", "Evaluation manifest"],
        ["Cache original inference once per model/sample.", "Generate or load the exact transformed input and run transformed inference.", "Compute class consistency, confidence_delta, and absolute_confidence_delta.", "Retain all valid prediction pairs for RQ1; skip CAM generation when class changes and record prediction_inconsistent.", "For consistent pairs, request paired CAMs using original_predicted_class, align, validate, and compute metrics.", "Reconcile expected pairs against valid, excluded, and failed outcomes."],
        ["Missing original/transformed pair", "Duplicate compound key", "Incompatible configuration/run identity", "Service failure without reason", "Population reconciliation mismatch"],
        ["A mock test proves XAIGenerator is never called for inconsistent pairs.", "Every expected pair reaches a valid or documented terminal state.", "JointEvaluator contains orchestration only, not scientific algorithms."],
    )

    module_spec(
        doc,
        "D10",
        "StatisticalAnalysis and RelationshipAnalyzer",
        "Estimate uncertainty, compare CAM methods on paired evidence, control multiplicity, and analyze RQ3 without invalid population mixing.",
        ["Validated per-sample JointRecords", "leaf_id grouping", "Predeclared statistics configuration", "Comparison-family definitions"],
        ["Descriptive summaries", "Cluster-bootstrap confidence intervals", "Paired test results", "Rank-biserial effect sizes", "Holm-adjusted p-values", "RQ3 relationship records"],
        ["Reconcile eligible populations and report image/leaf/exclusion counts.", "Resample leaf_id clusters with replacement and recompute declared statistics.", "Intersect exact paired keys before Wilcoxon comparisons.", "Calculate rank-biserial effect size and declared difference estimates.", "Apply Holm correction within predeclared families.", "For RQ3, analyze confidence change versus stability only in the consistent valid population and avoid causal claims."],
        ["Unpaired method samples", "Insufficient leaves/pairs", "Undefined all-zero-difference policy", "Missing comparison-family ID", "Attempt to assign stability to inconsistent samples"],
        ["Reference datasets reproduce expected results.", "Reports include n_images, n_leaves, n_pairs, exclusions, estimates, CI, raw/adjusted p-values, and effect size.", "No test is selected after observing results."],
    )

    module_spec(
        doc,
        "D11",
        "ResultStore and PartitionManager",
        "Persist immutable evidence atomically and support interruption-safe resume without duplicate or mixed-protocol records.",
        ["Typed records", "RunContext identity", "Partition key", "Expected schema/version", "Artifact root"],
        ["Atomic Parquet/JSON partitions", "Checksums", "Completion markers", "Resume status", "Uniqueness audit"],
        ["Validate schema and compound-key uniqueness before write.", "Write to a temporary run-scoped path, fsync where supported, then atomically rename.", "Record checksum and completion marker only after validation.", "On resume, reuse only partitions whose run/config/protocol/schema/checksums match.", "Reject append into finalized or incompatible partitions."],
        ["Duplicate key", "Checksum mismatch", "Schema mismatch", "Partial write", "Foreign run identity", "Finalized artifact mutation"],
        ["Injected interruption leaves no valid-looking partial artifact.", "Resume creates no duplicates and does not rerun completed valid partitions.", "Every stored row is attributable to one RunContext."],
    )

    module_spec(
        doc,
        "D12",
        "ArtifactGenerator",
        "Convert validated records into publication tables, figures, and reports without recomputing scientific metrics.",
        ["Validated summary/statistical records", "Visualization configuration", "Predeclared qualitative-selection rule", "RunContext"],
        ["CSV/Parquet", "LaTeX tables", "PNG 300 DPI and vector PDF figures", "Statistical report", "Qualitative selection manifest"],
        ["Read validated records only and verify source hashes.", "Generate CSV and LaTeX from the same data object and rounding policy.", "Generate figures with consistent axes, scales, colormaps, labels, CI notation, and sample counts.", "Select qualitative examples by a predeclared deterministic rule such as stability quantiles.", "Register every artifact and its source lineage."],
        ["Source hash mismatch", "Manual numeric override", "Unreconciled population", "Missing sample count/uncertainty", "Non-deterministic qualitative selection"],
        ["Tables and figures regenerate from per-sample evidence.", "CSV, LaTeX, and displayed numbers agree.", "No scientific metric implementation exists in this module."],
    )

    module_spec(
        doc,
        "D13",
        "RunContext, RunManifestWriter, and Provenance",
        "Bind all logs, records, commands, environments, configurations, and artifacts to one collision-resistant run identity.",
        ["Protocol/config hashes", "Dataset/manifest/split/checkpoint identities", "Git state", "Runtime environment", "User-approved run mode"],
        ["RunContext", "run_manifest.json", "artifact_index.json", "SHA256SUMS", "Reproducibility report", "Structured logs"],
        ["Create or inject a collision-resistant run_id and bind structured logging to it.", "Capture protocol, source, configuration, Git, environment, command, seed, hardware, and timestamp evidence.", "Refuse official execution when required clean-state/hash policies fail.", "Register artifacts and exclusions incrementally; finalize the manifest immutably.", "Support a clean rerun comparison using exact hashes for deterministic artifacts and declared tolerances for numerical outputs."],
        ["Arbitrary logging run_id not bound to RunContext", "Dirty official worktree when prohibited", "Hash or environment mismatch", "Manifest mutation after finalization", "Secret/local-only data detected"],
        ["Independent contexts do not leak identity or logging fields.", "Manifest JSON is complete and deterministic where required.", "Every final artifact has source lineage and SHA-256."],
    )


def build_part_e(doc: Document) -> None:
    major_section(doc, "PART E - DATA CONTRACTS")
    add_text(doc, "All scientific records are typed, schema-versioned, keyed explicitly, and immutable after finalization. Row order and batch position never carry identity.", italic=True)

    heading(doc, "E1. Contract Conventions", 2)
    add_bullets(doc, [
        "Identifiers are stable strings; timestamps use ISO 8601 with timezone; hashes use lowercase hexadecimal SHA-256.",
        "Paths stored in portable artifacts are canonical relative paths or approved artifact references, never machine-specific absolute paths.",
        "Enumerated status and reason values are validated; unknown values fail closed.",
        "Every record carries schema_version and run_id where generated during a run.",
        "Scientific records are written before aggregation and cannot be silently updated by visualization code.",
    ])

    heading(doc, "E2. SampleRecord", 2)
    contract_table(doc, [
        ["sample_id", "str", "YES", "Stable image identity derived from canonical relative path plus canonical RGB SHA-256; exact algorithm frozen in protocol"],
        ["leaf_id", "str", "YES", "Source-evidenced leaf grouping identity; may not be invented"],
        ["class_id", "int", "YES", "Frozen class index"],
        ["class_name", "str", "YES", "Canonical source class label"],
        ["source_split", "str", "YES", "Original dataset split"],
        ["split", "enum", "YES", "train | validation | test"],
        ["canonical_relative_path", "str", "YES", "Portable source path"],
        ["canonical_rgb_sha256", "str", "YES", "SHA-256 over canonical RGB representation"],
        ["width", "int", "YES", "Positive source image width"],
        ["height", "int", "YES", "Positive source image height"],
    ])

    heading(doc, "E3. TransformationRecord", 2)
    contract_table(doc, [
        ["run_id", "str", "YES", "Bound RunContext identity"],
        ["sample_id", "str", "YES", "References SampleRecord"],
        ["scenario_id", "str", "YES", "Unique configured transformation/severity identity"],
        ["transformation", "enum", "YES", "rotation | brightness | gaussian_noise | gaussian_blur"],
        ["severity", "enum", "YES", "mild | moderate | severe"],
        ["derived_seed", "int", "YES", "Stable per-sample seed"],
        ["parameters", "JSON", "YES", "Exact applied values, direction, kernel, sigma, factor, or angle"],
        ["interpolation", "str/null", "WHEN USED", "Frozen image interpolation policy"],
        ["border_policy", "str/null", "WHEN USED", "Rotation fill/border policy"],
        ["inverse_metadata", "JSON/null", "GEOMETRIC", "Exact inverse transform specification"],
        ["transformed_checksum", "str", "YES", "Checksum of deterministic transformed representation or cache artifact"],
    ])

    heading(doc, "E4. PredictionRecord", 2)
    contract_table(doc, [
        ["run_id", "str", "YES", "Bound run identity"],
        ["model_id", "str", "YES", "Backbone plus model configuration identity"],
        ["checkpoint_id", "str", "YES", "Approved checkpoint identity"],
        ["checkpoint_sha256", "str", "YES", "Verified checkpoint hash"],
        ["sample_id", "str", "YES", "Image identity"],
        ["scenario_id", "str", "YES", "original or transformed scenario"],
        ["true_class_id", "int", "YES", "Frozen label for descriptive correctness analysis"],
        ["predicted_class_id", "int", "YES", "Argmax predicted class"],
        ["confidence", "float", "YES", "Maximum softmax probability in [0,1]"],
        ["is_correct", "bool", "YES", "Predicted class equals true class"],
        ["logits_ref", "artifact ref/null", "OPTIONAL", "Reference to persisted logits where enabled"],
    ])

    heading(doc, "E5. PredictionPairRecord", 2)
    contract_table(doc, [
        ["run_id/model_id/sample_id/scenario_id", "compound key", "YES", "Unique prediction-pair identity"],
        ["original_predicted_class", "int", "YES", "Original decision target"],
        ["transformed_predicted_class", "int", "YES", "Decision after controlled transformation"],
        ["original_confidence", "float", "YES", "Original confidence"],
        ["transformed_confidence", "float", "YES", "Transformed confidence"],
        ["is_consistent", "bool", "YES", "Exact class-index equality"],
        ["confidence_delta", "float", "YES", "transformed minus original confidence"],
        ["absolute_confidence_delta", "float", "YES", "Absolute confidence difference"],
        ["eligible_for_explanation", "bool", "YES", "True only when class is consistent and required inference evidence exists"],
    ])

    heading(doc, "E6. CAMRecord", 2)
    contract_table(doc, [
        ["run_id/model_id/sample_id/scenario_id/xai_method", "compound key", "YES", "Unique CAM identity"],
        ["target_class", "int", "YES", "Must equal original predicted class for paired stability evaluation"],
        ["target_layer", "str", "YES", "Runtime-validated layer identifier"],
        ["heatmap_ref", "artifact ref", "YES", "Run-scoped array artifact or approved encoded representation"],
        ["heatmap_sha256", "str", "YES", "Checksum of stored heatmap"],
        ["height/width", "int", "YES", "Positive heatmap dimensions"],
        ["minimum/maximum", "float", "YES", "Raw or declared-stage extrema"],
        ["finite", "bool", "YES", "No NaN/Inf"],
        ["constant", "bool", "YES", "Constant/near-constant policy result"],
        ["generation_seconds", "float", "YES", "Measured runtime"],
        ["quality_status", "enum", "YES", "valid | invalid | failed"],
    ])

    heading(doc, "E7. StabilityRecord", 2)
    contract_table(doc, [
        ["run_id/model_id/sample_id/scenario_id/xai_method", "compound key", "YES", "Unique per-sample stability record"],
        ["leaf_id", "str", "YES", "Cluster identity for uncertainty"],
        ["transformation/severity", "enum", "YES", "Grouping factors"],
        ["class_name", "str", "YES", "Frozen class label"],
        ["target_class", "int", "YES", "Common decision target"],
        ["aligned", "bool", "YES", "Alignment completion"],
        ["valid_mask_pixels", "int", "YES", "Positive eligible pixel count"],
        ["ssim", "float/null", "CONDITIONAL", "SSIM with data_range=1.0"],
        ["pearson", "float/null", "CONDITIONAL", "Pearson correlation; null if invalid variance"],
        ["cosine", "float/null", "CONDITIONAL", "Cosine similarity; null if invalid norm"],
        ["included", "bool", "YES", "Included in official stability population"],
        ["exclusion_reason", "enum/null", "IF EXCLUDED", "Mandatory terminal reason when not included"],
    ])

    heading(doc, "E8. JointRecord and ExclusionRecord", 2)
    contract_table(doc, [
        ["joint key", "compound", "YES", "run_id + model_id + sample_id + scenario_id + xai_method"],
        ["prediction fields", "embedded/ref", "YES", "Consistency and confidence-change evidence"],
        ["stability fields", "embedded/ref/null", "CONDITIONAL", "Only when eligible and valid"],
        ["population", "enum", "YES", "prediction | explanation_eligible | valid_stability"],
        ["terminal_status", "enum", "YES", "valid | excluded | failed"],
        ["reason_code", "enum/null", "IF NOT VALID", "Machine-readable reason"],
        ["reason_detail", "str/null", "OPTIONAL", "Non-sensitive diagnostic context"],
        ["stage", "str", "YES", "Stage producing the terminal outcome"],
    ])

    heading(doc, "E9. RunManifest", 2)
    contract_table(doc, [
        ["run_id", "str", "YES", "Collision-resistant primary identity"],
        ["protocol_version/hash", "str", "YES", "Exact scientific authority"],
        ["config_bundle_sha256", "str", "YES", "Canonical resolved configuration"],
        ["dataset_revision/fingerprint", "str", "YES", "Immutable source identity"],
        ["manifest_hash/split_hashes", "str/map", "YES", "Data membership identity"],
        ["checkpoint_hashes", "map", "YES", "Approved model artifacts"],
        ["git_sha/git_dirty", "str/bool", "YES", "Code state"],
        ["python/torch/cuda/hardware", "JSON", "YES", "Runtime environment"],
        ["seed and command", "int/array", "YES", "Execution reproduction inputs"],
        ["timestamps/status", "JSON", "YES", "Lifecycle state"],
        ["artifacts/exclusions", "arrays", "YES", "Lineage and population accounting"],
    ])

    heading(doc, "E10. Formats, Naming, and Immutability", 2)
    add_table(
        doc,
        ["Artifact", "Format", "Naming / Location", "Immutability"],
        [
            ["Dataset manifest", "Parquet", "data/manifests/dataset_manifest.parquet", "Immutable after approval"],
            ["Frozen splits", "CSV/Parquet", "data/splits/<split>.*", "Immutable after freeze"],
            ["Per-sample records", "Parquet", "outputs/runs/<run_id>/<domain>/", "Append only through atomic partitions until finalization"],
            ["Configuration", "YAML", "outputs/runs/<run_id>/resolved_config.yaml", "Immutable for the run"],
            ["Run/artifact manifests", "JSON", "outputs/runs/<run_id>/", "Finalized immutably"],
            ["Tables", "CSV + LaTeX", "outputs/runs/<run_id>/tables/", "Regenerable; not source data"],
            ["Figures", "PNG + PDF", "outputs/runs/<run_id>/figures/", "Regenerable; not source data"],
        ],
        [1850, 1300, 3300, 2190],
        font_size=8.1,
    )


def build_part_f(doc: Document) -> None:
    major_section(doc, "PART F - BUSINESS RULES")
    callout(doc, "ENFORCEMENT", "Business rules are implementation-level expressions of the scientific protocol. A violation blocks the affected quality gate and must not be downgraded to a warning for an official run.", LIGHT_RED)

    rows = [
        ["BR-001", "Explanation stability is evaluated only when original and transformed predicted class indices are identical.", "JointEvaluator", "T-FIL-01"],
        ["BR-002", "Both CAMs in a stability pair use original_predicted_class as the target; adapters may not choose targets internally.", "JointEvaluator / XAI", "T-TGT-01"],
        ["BR-003", "Valid heatmaps are min-max normalized to [0,1] with epsilon 1e-8 before similarity calculation.", "HeatmapQuality", "T-XAI-01"],
        ["BR-004", "Constant, near-constant, zero-norm, non-finite, or otherwise invalid heatmaps receive no fabricated similarity value.", "HeatmapQuality / Metrics", "T-XAI-02"],
        ["BR-005", "All scientific parameters originate from versioned validated configuration; no hard-coded severity, target, metric, or statistical values.", "Config / all modules", "T-CFG-01"],
        ["BR-006", "Each transformation is deterministic from the frozen algorithm, global seed, sample_id, and scenario_id, and records exact applied parameters.", "Transformations", "T-TRF-01"],
        ["BR-007", "Geometric heatmaps are inverse-aligned using transformation metadata and compared only on a non-empty valid-overlap mask.", "Aligner", "T-ALN-01"],
        ["BR-008", "sample_id is stable and independent of row order, batch order, and absolute machine paths.", "Data", "T-DAT-01"],
        ["BR-009", "leaf_id must be source-evidenced; missing leaf grouping blocks the relevant official split/statistical design and is never silently fabricated.", "Data / Statistics", "T-DAT-02"],
        ["BR-010", "Official test data is not used for class selection, split tuning, severity pilot, hyperparameter tuning, checkpoint selection, or target-layer selection.", "Training / Governance", "T-TRN-01"],
        ["BR-011", "The official test population is preserved according to the approved source policy; an example ratio cannot override it.", "SplitManager", "T-SPL-01"],
        ["BR-012", "No sample_id or leaf_id may appear in more than one frozen split.", "SplitManager", "T-SPL-02"],
        ["BR-013", "Direction/sign selection for rotation and brightness is deterministic and recorded per sample.", "Transformations", "T-TRF-02"],
        ["BR-014", "Transformations operate on canonical pixels in [0,1] before model normalization.", "Data / Transformations", "T-TRF-03"],
        ["BR-015", "Both models receive the same transformed sample for the same sample_id and scenario_id.", "JointEvaluator", "T-TRF-04"],
        ["BR-016", "Original and transformed inference use the same approved checkpoint; all joins use explicit compound keys.", "Inference / Evaluation", "T-PAIR-01"],
        ["BR-017", "Paired statistical comparisons use the exact intersection of valid paired keys; tests are not run on aggregated group means.", "Statistics", "T-STA-01"],
        ["BR-018", "Cluster-aware uncertainty uses the approved leaf_id unit when multiple images per leaf exist.", "Statistics", "T-STA-02"],
        ["BR-019", "Holm correction is applied to all tests in each comparison family defined before result inspection.", "Statistics", "T-STA-03"],
        ["BR-020", "Every expected pair ends as valid, excluded, or failed with a machine-readable reason; no silent row loss is permitted.", "ResultStore / Evaluation", "T-REC-01"],
        ["BR-021", "Every change to scientific configuration creates a new configuration hash and a new run identity; incompatible outputs are never appended.", "Provenance", "T-PRO-01"],
        ["BR-022", "Writes are atomic and unique by compound key; resume reuses only complete checksum-valid partitions.", "ResultStore", "T-IO-01"],
        ["BR-023", "Publication tables and figures are generated from validated records and may not recompute or manually override scientific values.", "Artifacts", "T-ART-01"],
        ["BR-024", "Official runs obey clean-worktree, protocol-hash, config-hash, manifest-hash, split-hash, and checkpoint-hash policies.", "Provenance / CLI", "T-RUN-01"],
        ["BR-025", "Weights, raw data, secrets, credentials, local AI configuration, and absolute development paths are not committed or released.", "Repository / Release", "T-SEC-01"],
    ]
    add_table(doc, ["ID", "Rule", "Owner", "Test"], rows, [850, 4780, 1950, 1060], font_size=7.5)

    heading(doc, "F2. Prohibited Anti-Patterns", 2)
    add_bullets(doc, [
        "Joining results by DataFrame row number, filesystem order, or batch position.",
        "Using ground-truth class for one CAM and predicted class for another without a separately approved sensitivity analysis.",
        "Dropping NaN, failed CAMs, inconsistent predictions, or missing pairs without a reason record and denominator reconciliation.",
        "Normalizing a constant heatmap to zeros and allowing downstream metrics to treat zero-zero as perfect stability.",
        "Selecting severity, target layer, checkpoint, statistical test, or comparison family after viewing official test results.",
        "Copying numbers into manuscript tables or editing figures independently of validated source records.",
        "Reusing an existing run directory after scientific configuration, data membership, checkpoint, or code identity changes.",
        "Using a generic random seed call without deterministic worker and per-sample derivation policy.",
    ])


def build_part_g(doc: Document) -> None:
    major_section(doc, "PART G - SCIENTIFIC INVARIANTS")
    add_text(doc, "Scientific invariants protect validity rather than implementation convenience. Each invariant requires an executable test or an auditable gate before official evidence is accepted.", italic=True)
    add_table(
        doc,
        ["ID", "Invariant", "Verification"],
        [
            ["SI-01", "The same approved input, configuration, seed, model, and deterministic runtime path reproduce the same identity, transformation, filtering, and aggregation decisions.", "Determinism and clean-rerun tests"],
            ["SI-02", "No official test observation influences class, split, severity, training, checkpoint, target-layer, metric, or statistical-plan selection.", "Access instrumentation, configuration audit, and governance review"],
            ["SI-03", "No sample_id or leaf_id crosses frozen split boundaries.", "Set-intersection and duplicate-leakage tests"],
            ["SI-04", "A valid normalized heatmap lies in [0,1] and contains finite values.", "Assertions and property tests"],
            ["SI-05", "A constant/near-constant heatmap is invalid for official similarity and has a documented exclusion.", "Scientific negative test"],
            ["SI-06", "Explanation stability exists only for prediction-consistent pairs with the same explicit target class.", "Mock call-count and target-propagation tests"],
            ["SI-07", "For rotation, inverse alignment plus valid masking precedes heatmap comparison.", "Synthetic hotspot/alignment test"],
            ["SI-08", "SSIM(x,x), Pearson(x,x), and cosine(x,x) are approximately 1 for valid non-constant reference maps.", "Metric sanity tests"],
            ["SI-09", "Pearson with zero variance and cosine with zero norm are invalid rather than coerced to a scientific value.", "Metric negative tests"],
            ["SI-10", "Per-sample records are preserved before descriptive or inferential aggregation.", "Schema and lineage audit"],
            ["SI-11", "Paired comparisons use identical paired keys for both methods and report the actual number of pairs.", "Pair-intersection tests"],
            ["SI-12", "Cluster bootstrap resamples the approved independent unit and includes all eligible observations within a sampled cluster.", "Reference bootstrap fixture"],
            ["SI-13", "Holm-adjusted p-values are not smaller than corresponding raw p-values and obey the step-down monotonicity policy.", "Reference correction test"],
            ["SI-14", "Every expected evaluation pair is reconciled as valid, excluded, or failed with a reason.", "Population accounting assertion"],
            ["SI-15", "A finalized record, partition, manifest, or publication artifact is immutable for its run identity.", "Checksum and mutation tests"],
            ["SI-16", "Tables and figures are reproducible from validated machine-readable evidence without manual numeric edits.", "Artifact regeneration and data-diff test"],
            ["SI-17", "RQ3 does not assign an explanation-stability value to prediction-inconsistent samples and does not claim causality from association.", "Analysis-contract and report-lint tests"],
            ["SI-18", "Claims remain within the approved PlantVillage population and controlled-transformation design.", "Claim-boundary review checklist"],
        ],
        [900, 5280, 2460],
        font_size=7.9,
    )

    heading(doc, "G2. Reference Equations", 2)
    code_block(doc, "is_consistent = (original_predicted_class == transformed_predicted_class)\n\nconfidence_delta = transformed_confidence - original_confidence\nabsolute_confidence_delta = abs(confidence_delta)\n\nH_norm = (H - min(H)) / (max(H) - min(H) + 1e-8)\n\npearson(x,y) = cov(x,y) / (std(x) * std(y))\ncosine(x,y) = dot(x,y) / (norm(x) * norm(y))")


def build_part_h(doc: Document) -> None:
    major_section(doc, "PART H - QUALITY GATES")
    add_text(doc, "Completion of code is not authorization to proceed. Every gate requires machine-readable evidence and, where specified, human approval.", italic=True)
    add_table(
        doc,
        ["Gate", "Scope", "PASS Evidence", "Current State"],
        [
            ["G0A", "Software bootstrap readiness", "Package/tooling, draft protocol validation, logging/run identity foundation, tests, and environment checks meet approved bootstrap criteria", "PASS per project context"],
            ["G0B", "Scientific protocol freeze readiness", "Dataset audit, manifest, selected classes, severity pilot, selected checkpoints, target-layer mapping, canonical protocol v1, hash, and human approval", "BLOCKED"],
            ["G1", "Dataset and splits", "Immutable receipt/revision, canonical manifest, five approved classes, no corruption/conflicting duplicates/leakage, frozen split hashes", "PENDING"],
            ["G2", "Model baselines", "Smoke/tiny-batch checks, training evidence, validation-selected checkpoints, checkpoint hashes, baseline plan", "PENDING"],
            ["G3", "Transformations", "Twelve deterministic scenarios, exact parameter records, synthetic tests, semantic pilot approval", "PENDING"],
            ["G4", "XAI", "Three adapters, explicit target contract, runtime target-layer validation, finite non-trivial smoke heatmaps, hook cleanup", "PENDING"],
            ["G5", "Alignment and metrics", "Inverse-alignment reference tests, valid masks, normalization/invalid-map tests, metric reference match", "PENDING"],
            ["G6", "Joint evaluation", "Explicit unique keys, consistency call-path test, population reconciliation, atomic partitions, resume test", "PENDING"],
            ["G7", "Statistics", "Leaf-aware bootstrap reference, exact paired intersections, Wilcoxon/effect-size references, frozen families, Holm tests", "PENDING"],
            ["G8", "Publication artifacts", "All tables/figures regenerated from validated records, source hashes and sample counts recorded, visual/content audit", "PENDING"],
            ["G9", "Reproducibility release", "Fresh-environment rerun, exact deterministic hashes, numerical tolerance report, artifact index, privacy/secret scan", "PENDING"],
        ],
        [850, 2100, 4150, 1540],
        font_size=7.8,
    )

    heading(doc, "H2. G0B Blocking Evidence Register", 2)
    add_table(
        doc,
        ["Blocker", "Required Evidence", "Resolution Authority"],
        [
            ["Dataset audit not executed", "Immutable receipt, schema, integrity, duplicate, class, and grouping reports", "Research lead + data auditor"],
            ["Audited manifest not built", "Canonical manifest, schema version, deterministic rebuild, and hash", "Research lead"],
            ["Selected classes not recorded", "Class-selection decision with image/leaf counts and confound rationale", "Research lead"],
            ["Transformation pilot not executed", "Train/validation semantic and quality pilot", "Research lead + domain reviewer"],
            ["Selected severities not recorded", "Approved parameter values and direction/border policies", "Research lead"],
            ["Checkpoints not selected", "Validation-based evidence and artifact hashes", "Research lead + model reviewer"],
            ["Target layers not mapped", "Runtime activation/gradient/hook evidence for each backbone", "Research lead + XAI reviewer"],
        ],
        [2480, 4310, 1850],
        font_size=8.3,
    )

    heading(doc, "H3. Definition of Done for an Implementation Increment", 2)
    add_bullets(doc, [
        "Scope matches an approved task and does not introduce unapproved scientific execution.",
        "Public interfaces and data contracts are typed; configuration and failures validate explicitly.",
        "Unit, negative, integration, and relevant scientific tests pass.",
        "Formatting, linting, type checking, dependency integrity, and whitespace checks pass.",
        "No hard-coded scientific values, absolute paths, credentials, local AI artifacts, or generated experimental results are committed.",
        "Decision records and traceability are updated when an approved boundary changes.",
        "The reviewer can reproduce the evidence and map it to BR/SI/test identifiers.",
    ])

    heading(doc, "H4. Release Stop Conditions", 2)
    add_bullets(doc, [
        "Unreconciled expected pairs, missing reason codes, duplicate keys, or silent row loss.",
        "Protocol, configuration, manifest, split, checkpoint, or Git identity mismatch.",
        "Official results produced while G0B is blocked or protocol is not frozen.",
        "Manual edits to computed evidence or inconsistent values across CSV, LaTeX, figures, and reports.",
        "Unresolved high-severity scientific validity finding or failed reproducibility comparison.",
    ])


def build_part_i(doc: Document) -> None:
    major_section(doc, "PART I - TESTING STRATEGY")
    add_text(doc, "Testing is organized by failure cost: broad unit coverage, cross-module integration tests, focused scientific invariant tests, and clean-environment reproducibility tests.", italic=True)

    heading(doc, "I1. Test Pyramid", 2)
    code_block(doc, "            Reproducibility / Release Tests\n                 Scientific Tests\n              Integration / Contract Tests\n        Unit / Property / Negative Tests\n   Static Analysis, Schema Validation, and Scans")
    add_table(
        doc,
        ["Category", "Purpose", "Typical Scope"],
        [
            ["Static and schema", "Prevent malformed code/configuration and local/privacy leakage", "Ruff, MyPy, schema validation, dependency checks, secret/local-only scans"],
            ["Unit/property/negative", "Verify functions, boundary conditions, and fail-closed behavior", "Identity, hashing, transforms, normalization, metrics, pairing, correction, atomic I/O"],
            ["Integration/contract", "Verify module interactions and record contracts", "Data-to-transform-to-inference; model-to-CAM-to-alignment-to-metric; records-to-statistics"],
            ["Scientific", "Protect research invariants and population rules", "Consistency gate, same target, leakage, invalid heatmaps, alignment, leaf clustering"],
            ["Reproducibility", "Verify a clean rerun reconstructs the same evidence", "Environment, manifests, partitions, artifact regeneration, numerical tolerance"],
        ],
        [1800, 3100, 3740],
        font_size=8.3,
    )

    heading(doc, "I2. Mandatory Test Catalogue", 2)
    add_table(
        doc,
        ["Test ID", "Test", "Expected Result", "Level"],
        [
            ["T-CFG-01", "Unknown key or invalid cross-field protocol choice", "Validation fails with actionable path; no silent fallback", "Unit/Scientific"],
            ["T-DAT-01", "Build sample IDs twice with shuffled row order", "Same per-image IDs and manifest hash", "Property"],
            ["T-DAT-02", "Source lacks trustworthy leaf_id", "Official grouping/split/statistics gate fails", "Scientific negative"],
            ["T-SPL-01", "Official-test policy versus configured ratio", "Approved source policy wins or protocol blocks", "Scientific"],
            ["T-SPL-02", "Inject sample/leaf overlap across splits", "G1 fails and reports exact overlaps", "Integration"],
            ["T-TRF-01", "Repeat each scenario with same sample/seed", "Pixel-identical transformed input and metadata", "Property"],
            ["T-TRF-02", "Rotation/brightness sign derivation", "Stable recorded direction for each pair", "Unit"],
            ["T-TRF-03", "Apply transform at wrong preprocessing stage", "Contract/assertion rejects normalized model tensor", "Negative"],
            ["T-TRF-04", "Two models request same sample/scenario", "Both receive identical transformed checksum", "Integration"],
            ["T-TRN-01", "Instrument data access during model selection", "No official test access occurs", "Scientific"],
            ["T-TGT-01", "Generate paired CAMs", "Both adapters receive original predicted class", "Integration/Scientific"],
            ["T-XAI-01", "Normalize finite non-constant heatmap", "Range is [0,1] with expected reference values", "Unit"],
            ["T-XAI-02", "Constant/NaN/Inf/zero-norm heatmaps", "Invalid status and reason; no official metric", "Scientific negative"],
            ["T-ALN-01", "Rotate synthetic hotspot and inverse-align", "Post-alignment similarity exceeds pre-alignment and reference threshold", "Scientific"],
            ["T-PAIR-01", "Original/transformed predictions use different checkpoint or row join", "Contract rejects mismatch; explicit keys required", "Integration negative"],
            ["T-FIL-01", "Prediction class changes", "Prediction record retained; XAIGenerator call count is zero", "Scientific integration"],
            ["T-REC-01", "Expected population with injected failures", "Valid + excluded + failed equals expected, with reasons", "Scientific"],
            ["T-STA-01", "CAM method sets have unequal membership", "Exact paired intersection and n_pairs are reported", "Unit/Scientific"],
            ["T-STA-02", "Multiple images per leaf", "Bootstrap samples leaf clusters, not independent images", "Reference statistical"],
            ["T-STA-03", "Known p-value family", "Holm adjusted values match reference and monotonicity rules", "Unit"],
            ["T-IO-01", "Interrupt atomic partition write", "No complete marker or readable official partial artifact remains", "Integration"],
            ["T-PRO-01", "Change scientific config under existing run identity", "Write/resume is rejected; new run required", "Integration"],
            ["T-ART-01", "Regenerate CSV/LaTeX/figures", "Displayed values and source hashes agree with validated records", "Reproducibility"],
            ["T-RUN-01", "Dirty or hash-mismatched official run", "Execution fails closed before scientific work", "Operational"],
            ["T-SEC-01", "Release scan", "No secrets, raw data, weights, local AI configuration, or absolute local paths", "Release"],
        ],
        [1050, 3260, 3270, 1060],
        font_size=7.35,
    )

    heading(doc, "I3. Test Fixtures and Golden References", 2)
    add_bullets(doc, [
        "Tiny deterministic RGB images with known hashes and labels.",
        "Synthetic leaf groups that expose image-level versus cluster-level leakage.",
        "Synthetic hotspot heatmaps for rotation and mask validation.",
        "Constant, near-constant, zero-norm, NaN, Inf, and mismatched-shape heatmaps.",
        "Reference arrays for SSIM, Pearson, cosine, bootstrap, Wilcoxon, rank-biserial, and Holm calculations.",
        "Interrupted and duplicate partition fixtures for atomic/resume tests.",
        "Minimal validated per-sample dataset capable of regenerating golden tables and figures without claiming scientific results.",
    ])

    heading(doc, "I4. Determinism and Numerical Tolerance", 2)
    add_table(
        doc,
        ["Class", "Expected Guarantee", "Examples"],
        [
            ["Exact determinism", "Byte/identity exact where practical", "Sample IDs, manifest/split membership, seeds, transform parameters, filtering, keys, artifact selection"],
            ["Numerical reproducibility", "Within a predeclared absolute/relative tolerance", "GPU training, CAM arrays, floating-point summaries"],
            ["Statistical reproducibility", "Same algorithm, seed, clusters, pair keys, family definitions, and reference tolerance", "Bootstrap CI, Wilcoxon, effect sizes, corrected p-values"],
        ],
        [1800, 3450, 3390],
        font_size=8.4,
    )

    heading(doc, "I5. Coverage and Evidence Policy", 2)
    add_bullets(doc, [
        "Coverage percentage is a supporting signal, not a substitute for scientific tests; the project target is greater than 80% unit coverage for eligible implementation modules.",
        "Every BR and SI must map to at least one named test or manual gate evidence item.",
        "Failed tests are reported; they are not bypassed by regenerating fixtures from the failing output.",
        "Official gate evidence records command, environment, Git SHA, result counts, failures, warnings, and artifact hashes.",
    ])


def build_part_j(doc: Document) -> None:
    major_section(doc, "PART J - HANDOFF PROTOCOL")
    add_text(doc, "Handoff converts the specification into bounded implementation work while preserving human authority over scientific choices.", italic=True)

    heading(doc, "J1. Handoff Package", 2)
    add_bullets(doc, [
        "This specification and the active draft/frozen protocol bundle.",
        "Approved decision records for dataset source, classes, split policy, severity, checkpoints, target layers, statistical units, and comparison families.",
        "Typed schemas and example non-scientific fixtures for every record contract.",
        "Reference implementations or validated formulas where available.",
        "Golden files for identity, transformation, alignment, metrics, statistics, and artifact generation.",
        "Gate evidence, unresolved decision register, and explicit authorization boundary.",
    ])

    heading(doc, "J2. Task Breakdown Template", 2)
    code_block(doc, "Task ID: TASK-<MODULE>-<N>\nModule: <package / class>\nPriority: P0 | P1 | P2\nMode: design | implementation | audit | remediation\nAuthorized files: <exact allowlist>\nDependencies: <task IDs / gate evidence>\nScientific execution authorized: YES | NO\n\nObjective:\n  <one bounded outcome>\n\nApplicable rules:\n  BR-xxx; SI-xxx; contracts; gate\n\nAcceptance criteria:\n  [ ] Typed interface and explicit validation\n  [ ] Positive, negative, integration, and scientific tests as applicable\n  [ ] No scientific values invented or moved outside configuration\n  [ ] Gate evidence and traceability updated\n\nProhibited actions:\n  <Git, protocol, dataset, network, training, or result actions outside scope>\n\nFiles to create/modify:\n  <exact paths>")

    heading(doc, "J3. AI Pair-Programming Guardrail", 2)
    code_block(doc, "ROLE\nYou are a Research Software Engineer implementing one approved task.\n\nAUTHORITATIVE CONTEXT\n- Project: PlantXAI-Stability\n- Protocol: <version/status/hash>\n- Gate state: <G0A/G0B/...>\n- Authorized files: <allowlist>\n- Applicable BR/SI/data contracts: <IDs>\n\nREQUIREMENTS\n- Preserve typed boundaries and fail-closed validation.\n- Add tests that prove scientific call paths, not only final outputs.\n- Record uncertainty or blockers; do not invent scientific values.\n\nPROHIBITED\n- Change/freeze protocol without human approval.\n- Use official test for selection.\n- Fabricate dataset counts, metrics, p-values, CIs, results, or claims.\n- Hard-code scientific parameters or absolute local paths.\n- Modify Git state or files beyond authorization.")

    heading(doc, "J4. Review Checklist", 2)
    add_bullets(doc, [
        "Implementation remains inside the task/file allowlist and does not expand scientific scope.",
        "Data contracts, compound keys, status enums, and exclusion reasons match this specification.",
        "Business rules and scientific invariants are enforced by tests or gate evidence.",
        "No test leakage, silent exclusions, manual result edits, or untracked scientific defaults exist.",
        "Runtime failures terminate in explicit records and do not corrupt prior partitions.",
        "Documentation distinguishes APPROVED, PROVISIONAL, PENDING, and BLOCKED values.",
        "Traceability and decision records are updated before merge or release.",
    ])

    heading(doc, "J5. Change-Control Protocol", 2)
    add_numbered(doc, [
        "Raise a Decision Record identifying the scientific or architectural issue, alternatives, impact, and evidence.",
        "Classify the change as implementation-only, protocol-draft change, or frozen-protocol change.",
        "Obtain human research-lead approval for every scientific change.",
        "Update configuration/schema/specification/tests and compute a new canonical hash.",
        "If the protocol was frozen or official evidence exists, create a new protocol/run lineage; never overwrite the previous run.",
        "Regenerate affected gate evidence and record deviations in release documentation.",
    ])

    heading(doc, "J6. Responsibility Matrix", 2)
    add_table(
        doc,
        ["Role", "Accountability", "Cannot Delegate"],
        [
            ["Research lead", "Scientific questions, protocol choices, claim boundaries, freeze, and release authorization", "Final approval and interpretation"],
            ["Research software architect", "Module boundaries, contracts, dependency rules, gate design, and traceability", "Architectural integrity review"],
            ["Developer", "Implementation, tests, documentation, and evidence within authorized scope", "Reporting failures and uncertainties"],
            ["Independent reviewer/auditor", "Read-only verification of repository, records, gates, statistics, and reproducibility", "Independent gate conclusion"],
            ["AI pair programmer", "Bounded assistance under explicit prompt, allowlist, and rules", "Human scientific judgment or approval"],
        ],
        [1650, 4400, 2590],
        font_size=8.3,
    )

    heading(doc, "J7. Research Traceability Matrix", 2)
    add_table(
        doc,
        ["RQ", "Evidence Records", "Modules", "Key Rules / Tests", "Publication Artifacts"],
        [
            ["RQ1", "PredictionRecord; PredictionPairRecord; transformation metadata", "D4, D5, D9, D10", "BR-006, BR-010, BR-015/016; T-TRF-01, T-TRN-01, T-FIL-01", "Prediction summary; Table 3; consistency-vs-severity figure"],
            ["RQ2", "CAMRecord; StabilityRecord; ExclusionRecord", "D6-D10", "BR-001-004, BR-007, BR-017-020; T-TGT-01, T-XAI-02, T-ALN-01, T-STA-01/03", "Table 4; Table 5; stability figures; qualitative examples"],
            ["RQ3", "JointRecord with confidence and stability fields", "D9-D10", "SI-06, SI-14, SI-17; T-FIL-01, T-REC-01", "Relationship table/figure; stable-prediction/unstable-explanation cases"],
        ],
        [650, 2150, 1200, 2750, 1890],
        font_size=7.5,
    )

    heading(doc, "J8. Final Release Checklist", 2)
    add_bullets(doc, [
        "Protocol is approved, frozen, hash-verified, and linked to the run.",
        "G0-G9 evidence is complete and no unresolved stop condition remains.",
        "Expected populations reconcile and every exclusion/failure is documented.",
        "Statistical families, units, tests, effect sizes, CIs, and corrected values match the frozen plan.",
        "Tables, figures, and narrative claims are traceable and within scope.",
        "Clean rerun evidence, SHA256SUMS, artifact index, and environment record are present.",
        "Privacy, licensing, secrets, raw-data, weight, and local-only configuration checks pass.",
    ])


def build_appendices(doc: Document) -> None:
    major_section(doc, "APPENDIX A - PROVISIONAL CONFIGURATION BASELINE")
    callout(doc, "PROVISIONAL", "The following values are candidate defaults from the design materials. They are not authorized official parameters until pilot evidence and human approval are recorded in a frozen protocol.", LIGHT_AMBER)

    heading(doc, "A.1 Dataset and Preprocessing", 2)
    add_table(
        doc,
        ["Parameter", "Candidate Value", "Status / Freeze Requirement"],
        [
            ["Dataset source", "mohanty/PlantVillage; color configuration", "PENDING immutable operational and upstream revision"],
            ["Classes", "Five tomato classes listed in Part A", "PENDING audit and class-selection approval"],
            ["Image color", "Canonical RGB", "PLANNED"],
            ["Input size", "224 x 224", "PROVISIONAL model preprocessing"],
            ["ImageNet mean", "[0.485, 0.456, 0.406]", "PROVISIONAL; tie to selected pretrained weights"],
            ["ImageNet std", "[0.229, 0.224, 0.225]", "PROVISIONAL; tie to selected pretrained weights"],
            ["Split", "Preserve official test; split approved train into train/validation by leaf_id and class", "PENDING dataset-policy confirmation"],
        ],
        [2300, 3100, 3240],
    )

    heading(doc, "A.2 Transformation Scenarios", 2)
    add_table(
        doc,
        ["Transformation", "Mild", "Moderate", "Severe", "Additional Policy"],
        [
            ["Rotation", "+/-10 deg", "+/-25 deg", "+/-45 deg", "Bilinear; fill policy PENDING; inverse alignment and valid mask required"],
            ["Brightness", "0.90 / 1.10", "0.70 / 1.30", "0.50 / 1.50", "Direction deterministic; clip [0,1]"],
            ["Gaussian noise", "sigma 0.01", "sigma 0.05", "sigma 0.10", "Mean 0; deterministic per-sample noise; clip [0,1]"],
            ["Gaussian blur", "3x3; sigma 0.8", "5x5; sigma 1.5", "9x9; sigma 3.0", "Exact kernel/sigma implementation to be frozen"],
        ],
        [1650, 1300, 1300, 1300, 3090],
        font_size=8.0,
    )

    heading(doc, "A.3 Model Training", 2)
    add_table(
        doc,
        ["Parameter", "Candidate Value", "Status"],
        [
            ["Backbones", "ResNet50; EfficientNet-B0", "PROVISIONAL"],
            ["Pretraining", "ImageNet weights", "PROVISIONAL; exact weight enum/version required"],
            ["Classes", "5", "PENDING class approval"],
            ["Maximum epochs", "50", "PROVISIONAL"],
            ["Batch size", "32", "PROVISIONAL; resource dependent but recorded"],
            ["Learning rate", "0.001", "PROVISIONAL"],
            ["Optimizer", "AdamW", "PROVISIONAL"],
            ["Weight decay", "0.0001", "PROVISIONAL"],
            ["Scheduler", "Cosine", "PROVISIONAL exact schedule"],
            ["Mixed precision", "Enabled where supported", "Execution parameter; record exact scaler/runtime"],
            ["Early stopping patience", "8", "PROVISIONAL"],
            ["Selection metric", "Validation macro-F1", "PROVISIONAL; must be predeclared"],
            ["Global seed", "42", "PROVISIONAL"],
        ],
        [2750, 3000, 2890],
    )

    heading(doc, "A.4 XAI and Statistics", 2)
    add_table(
        doc,
        ["Parameter", "Candidate Value", "Status"],
        [
            ["ResNet50 target layer", "layer4.2.conv3", "PROVISIONAL; runtime validation required"],
            ["EfficientNet-B0 target layer", "features.8.0", "PROVISIONAL; runtime validation required"],
            ["XAI methods", "Grad-CAM; Grad-CAM++; Score-CAM", "PLANNED; implementation/library versions to be pinned"],
            ["XAI batch size", "1", "Execution default; record actual"],
            ["Heatmap epsilon", "1e-8", "PLANNED"],
            ["SSIM", "data_range = 1.0", "PLANNED; window/boundary policy to be frozen"],
            ["Bootstrap", "10,000 iterations; 95% CI; seed 42; leaf_id cluster", "PROVISIONAL method/percentile policy"],
            ["Paired test", "Wilcoxon signed-rank; Pratt zero handling", "PROVISIONAL exact library and alternative"],
            ["Effect size", "Rank-biserial correlation", "PLANNED implementation/reference"],
            ["Multiple testing", "Holm; family-wise alpha 0.05", "PROVISIONAL family definitions"],
        ],
        [2750, 3350, 2540],
        font_size=8.2,
    )

    major_section(doc, "APPENDIX B - UNRESOLVED DECISION REGISTER")
    add_table(
        doc,
        ["ID", "Decision Required", "Evidence / Resolution", "Blocks"],
        [
            ["PND-001", "Immutable dataset revision and source receipt", "Pin upstream/operational identifiers and verify schema", "G1/G0B"],
            ["PND-002", "Trustworthy leaf_id source and grouping policy", "Dataset evidence and duplicate/group audit", "G1/G7/G0B"],
            ["PND-003", "Five-class approval and effective sample size", "Image/leaf counts, confound analysis, decision record", "G1/G0B"],
            ["PND-004", "Official test versus 70/10/20 example", "One canonical split policy", "G1/G0B"],
            ["PND-005", "Severity values and semantic preservation", "Train/validation pilot and human review", "G3/G0B"],
            ["PND-006", "Rotation fill/interpolation/mask policy", "Synthetic alignment and artifact-risk evaluation", "G3/G5"],
            ["PND-007", "Training hyperparameter selection procedure", "Predeclared train/validation tuning policy", "G2"],
            ["PND-008", "Selected checkpoints", "Validation evidence and hashes", "G2/G0B"],
            ["PND-009", "Target-layer mapping", "Activation/gradient/hook/runtime evidence", "G4/G0B"],
            ["PND-010", "Constant/near-constant heatmap threshold", "Reference fixtures and method-independent quality policy", "G5"],
            ["PND-011", "Bootstrap CI method and primary statistic", "Statistical analysis decision record", "G7"],
            ["PND-012", "Wilcoxon alternative, zero policy, and comparison families", "Predeclared statistical plan", "G7"],
            ["PND-013", "RQ3 primary association measure", "Define continuous robustness quantity and correlation/model", "G7"],
            ["PND-014", "Numerical reproducibility tolerances", "Hardware/runtime pilot and release policy", "G9"],
        ],
        [1050, 3300, 3270, 1020],
        font_size=7.6,
    )

    major_section(doc, "APPENDIX C - GLOSSARY")
    add_table(
        doc,
        ["Term", "Definition"],
        [
            ["Prediction robustness", "Stability of predicted class and confidence under a declared controlled transformation."],
            ["Explanation stability", "Similarity of paired explanations for the same decision target after required spatial alignment and quality validation."],
            ["Faithfulness", "Whether an explanation accurately reflects the model mechanism; not established by stability alone."],
            ["Prediction-consistent pair", "An original/transformed pair whose predicted class indices are identical."],
            ["Valid-overlap mask", "Pixels representing shared valid spatial support after geometric transformation and inverse alignment."],
            ["Scientific invariant", "A property that must hold for the evidence to be considered scientifically valid."],
            ["Decision Record", "A versioned document stating a choice, alternatives, evidence, authority, and consequences."],
            ["RunContext", "The collision-resistant identity and provenance binding for one execution lineage."],
            ["Official run", "An authorized execution under a frozen, hash-verified protocol and complete gate evidence."],
            ["Exploratory analysis", "Clearly labeled analysis not used as confirmatory evidence and not allowed to rewrite the frozen plan."],
        ],
        [2200, 6440],
        font_size=8.7,
    )

    major_section(doc, "CONCLUSION")
    add_text(doc, "This specification is the scientific and engineering contract among the research lead, software architect, developers, reviewers, and automation tools. It defines not only what the software must compute, but also which populations may be compared, how evidence must be keyed and preserved, when execution must stop, and which claims the resulting artifacts may support.")
    callout(doc, "FINAL PRINCIPLE", "Controlled data -> reproducible transformations -> paired predictions -> consistency gate -> same-target explanations -> spatial alignment -> validated stability metrics -> paired leaf-aware statistics -> reproducible evidence and bounded claims.", LIGHT_GREEN)


def build_document() -> None:
    if file_sha256(REFERENCE) != REFERENCE_SHA256:
        raise RuntimeError("Retained template hash mismatch; fresh distillation is required")
    shutil.copyfile(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    clear_body(doc)
    configure_styles(doc)
    doc.core_properties.title = "PlantXAI-Stability Research Software Design & Specification"
    doc.core_properties.subject = "Scientific design, software architecture, data contracts, QA, and handoff"
    doc.core_properties.keywords = "PlantXAI-Stability, XAI, robustness, stability, reproducibility"
    add_title_page(doc)
    add_front_matter(doc)
    build_part_a(doc)
    build_part_b(doc)
    build_part_c(doc)
    build_part_d(doc)
    build_part_e(doc)
    build_part_f(doc)
    build_part_g(doc)
    build_part_h(doc)
    build_part_i(doc)
    build_part_j(doc)
    build_appendices(doc)
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "size": OUTPUT.stat().st_size, "sha256": file_sha256(OUTPUT)}, indent=2))


if __name__ == "__main__":
    build_document()
